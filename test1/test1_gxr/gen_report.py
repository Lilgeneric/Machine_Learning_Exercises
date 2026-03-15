from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

DOCX_PATH = '/home/vla/machine_learning/test1/test1_gxr/test1_Experiment_Report.docx'
IMG_CM    = '/home/vla/machine_learning/test1/test1_gxr/output/confusion_matrices.png'
IMG_ROC   = '/home/vla/machine_learning/test1/test1_gxr/output/roc_curves.png'
OUT_PATH  = '/home/vla/machine_learning/test1/test1_gxr/test1_Experiment_Report_filled.docx'

doc = Document(DOCX_PATH)
table = doc.tables[0]

# ─── 辅助函数 ────────────────────────────────────────────────────────────────
def fmt_run(run, name='宋体', size=12, bold=False, color=None, code=False):
    fn = 'Courier New' if code else name
    run.font.name = fn
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), name)
    rPr.insert(0, rFonts)

def cell_clear(cell):
    for p in cell.paragraphs[1:]:
        p._element.getparent().remove(p._element)
    cell.paragraphs[0].clear()

def add_para(cell, text, bold=False, code=False, indent_first=False,
             align=WD_ALIGN_PARAGRAPH.LEFT, space_after=2):
    p = cell.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after  = Pt(space_after)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    if indent_first:
        pf.first_line_indent = Pt(24)
    p.alignment = align
    run = p.add_run(text)
    if code:
        fmt_run(run, code=True, size=9, color=(0x1a, 0x53, 0x76))
    else:
        fmt_run(run, bold=bold, size=12)
    return p

def add_sub(cell, text):
    p = cell.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(5)
    pf.space_after  = Pt(1)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = p.add_run(text)
    fmt_run(run, bold=True, size=12)
    return p

def add_code(cell, text):
    p = cell.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after  = Pt(0)
    pf.left_indent  = Cm(0.5)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = p.add_run(text)
    fmt_run(run, code=True, size=9, color=(0x1a, 0x53, 0x76))
    return p

def add_img(cell, path, width_cm=14):
    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(path, width=Cm(width_cm))
    return p

def build_table(cell, headers, rows, note=None):
    t = cell.add_table(rows=1+len(rows)+(1 if note else 0), cols=len(headers))
    pass  # no Table Grid in this doc
    for ci, h in enumerate(headers):
        c = t.rows[0].cells[ci]
        c.paragraphs[0].clear()
        r = c.paragraphs[0].add_run(h)
        fmt_run(r, bold=True, size=10)
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for ri, rd in enumerate(rows):
        for ci, val in enumerate(rd):
            c = t.rows[ri+1].cells[ci]
            c.paragraphs[0].clear()
            r = c.paragraphs[0].add_run(str(val))
            fmt_run(r, size=10)
            c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if note:
        nr = t.rows[-1]
        nr.cells[0].paragraphs[0].clear()
        r = nr.cells[0].paragraphs[0].add_run(note)
        fmt_run(r, size=9)
        for i in range(1, len(headers)):
            nr.cells[0].merge(nr.cells[i])
    return t

# ══════════════════════════════════════════════════════════════════════════════
# 封面信息
# ══════════════════════════════════════════════════════════════════════════════
for cell in table.rows[0].cells:
    if '专业' in cell.text or '学号' in cell.text:
        cell_clear(cell)
        r = cell.paragraphs[0].add_run('专业：电子信息工程       学号：2023XXXXXXXXX')
        fmt_run(r, size=12)
        break
for cell in table.rows[0].cells:
    if cell.text.strip() == '姓名':
        cell_clear(cell)
        r = cell.paragraphs[0].add_run('姓名：gxr')
        fmt_run(r, size=12)
        break
for cell in table.rows[1].cells:
    if '实验题目' in cell.text:
        cell_clear(cell)
        r = cell.paragraphs[0].add_run('实验题目：基于逻辑回归的银行营销结果预测')
        fmt_run(r, bold=True, size=12)
        break
for cell in table.rows[2].cells:
    if '实验时间' in cell.text:
        cell_clear(cell)
        r = cell.paragraphs[0].add_run('实验时间：2026年3月15日')
        fmt_run(r, size=12)
        break

# ══════════════════════════════════════════════════════════════════════════════
# 一、实验目的
# ══════════════════════════════════════════════════════════════════════════════
c5 = table.rows[5].cells[0]
cell_clear(c5)
add_para(c5, '一、实验目的', bold=True)
goals = [
    '1. 掌握逻辑回归（Logistic Regression）模型的数学原理与工程实现，理解正则化参数 C、惩罚项 penalty 对模型复杂度和泛化能力的影响；',
    '2. 学会在真实银行营销数据集上完成端到端的机器学习流程：数据探索→预处理→特征工程→模型训练→超参数调优→多指标评估；',
    '3. 深入理解类别不平衡问题的本质与危害，掌握并对比三种主流应对策略（SMOTE过采样 / 随机欠采样 / 损失权重调节），能够根据业务场景做出合理选择；',
    '4. 掌握网格搜索（Grid Search）结合K折分层交叉验证（Stratified K-Fold CV）的超参数调优方法，理解为何使用 roc_auc 而非 accuracy 作为评分指标；',
    '5. 理解在类别不平衡场景下各评价指标的含义与局限：Precision（预测订阅中有多少真的订阅）、Recall（真实订阅者中有多少被发现）、F1-Score（综合衡量）、AUC（排序能力）；',
    '6. 通过多轮迭代特征工程，记录并分析 AUC 的提升路径，理解特征质量对线性模型性能上限的决定性影响。',
]
for g in goals:
    add_para(c5, g, indent_first=True)

# ══════════════════════════════════════════════════════════════════════════════
# 二、实验项目内容
# ══════════════════════════════════════════════════════════════════════════════
c6 = table.rows[6].cells[0]
cell_clear(c6)
add_para(c6, '二、实验项目内容', bold=True)

add_sub(c6, '（一）数据集简介')
add_para(c6, '使用 UCI Bank Marketing 数据集（bank-additional-full.csv），共 41,188 条记录、20个特征字段、1个目标变量 y（yes=已订阅定期存款，no=未订阅）。数据由葡萄牙某银行2008—2013年直销电话活动记录汇编而成。', indent_first=True)
add_para(c6, '特征分三类：①客户属性（age、job、marital、education、default、housing、loan）；②营销接触信息（contact、month、day_of_week、duration、campaign、pdays、previous、poutcome）；③宏观经济指标（emp.var.rate、cons.price.idx、cons.conf.idx、euribor3m、nr.employed）。', indent_first=True)
add_para(c6, '类别分布：未订阅 36,548 条（88.73%），已订阅 4,640 条（11.27%），正负样本比约 1:7.9，存在严重类别不平衡，是本实验的核心挑战。', indent_first=True)

add_sub(c6, '（二）实验核心任务')
tasks = [
    '① 数据预处理：unknown缺失处理、education有序标签编码、9列无序特征One-Hot编码、StandardScaler特征标准化（防数据泄露）；',
    '② 特征工程：新增 previously_contacted 二值标志、log_duration 对数变换、精选 Top-15 特征的二阶交互项，最终特征维度从原始58维扩展至178维；',
    '③ 类别不平衡处理：SMOTE（sampling_strategy=0.4）/ RandomUnderSampler / class_weight=balanced 三种策略，均仅作用于训练集；',
    '④ 超参数调优：GridSearchCV 在 C×penalty×solver 空间搜索，5折分层CV，roc_auc评分；',
    '⑤ 多维评估：classification_report 输出 Precision/Recall/F1，绘制混淆矩阵和ROC曲线，计算AUC；',
    '⑥ 迭代优化记录：系统对比从基础版（AUC≈0.944）到最终版（AUC≈0.950）的完整优化路径。',
]
for t in tasks:
    add_para(c6, t, indent_first=True)

# ══════════════════════════════════════════════════════════════════════════════
# 三、实验过程或算法
# ══════════════════════════════════════════════════════════════════════════════
c7 = table.rows[7].cells[0]
cell_clear(c7)
add_para(c7, '三、实验过程或算法', bold=True)

# 3.1
add_sub(c7, '3.1 数据预处理')
add_para(c7, '（1）unknown 缺失值处理', bold=True)
add_para(c7, '数据集中 job（330条）、marital（80条）、education（1731条）、default（8597条）、housing（990条）、loan（990条）含有 unknown 字符串，合计约 12,718 个。本实验选择将 unknown 保留为独立类别而非填充，原因：unknown 可能携带有效信息（如客户拒绝披露信息本身就是一种行为特征），强行以众数填充会引入虚假信息。经One-Hot编码后，每个含unknown的字段会新增 _unknown 列，模型自行学习其权重。', indent_first=True)

add_para(c7, '（2）有序特征 Label Encoding：education', bold=True)
add_para(c7, '学历具有明确的语义顺序，采用整数映射保留顺序信息，避免One-Hot编码将有序关系拆散：', indent_first=True)
add_code(c7, "edu_order = {'illiterate':0,'basic.4y':1,'basic.6y':2,'basic.9y':3,")
add_code(c7, "             'high.school':4,'professional.course':5,'university.degree':6,'unknown':3}")
add_code(c7, "df['education'] = df['education'].map(edu_order)")

add_para(c7, '（3）无序特征 One-Hot Encoding', bold=True)
add_para(c7, 'job、marital、default、housing、loan、contact、month、day_of_week、poutcome 共9列无内在大小关系，使用 pd.get_dummies(drop_first=False) 全量展开，保留所有 dummy 列（含 unknown），编码后基础特征共58维。', indent_first=True)

add_para(c7, '（4）数据集划分与特征缩放（防泄露规范）', bold=True)
add_para(c7, '先按 stratify=y 做 8:2 分层划分，再对训练集 fit StandardScaler，最后用 transform 分别处理训练集和测试集。严禁对完整数据集 fit，否则测试集的统计量会"泄露"到预处理阶段，导致评估偏乐观。', indent_first=True)
add_code(c7, "X_train, X_test, y_train, y_test = train_test_split(")
add_code(c7, "    X, y, test_size=0.2, random_state=42, stratify=y)")
add_code(c7, "scaler = StandardScaler()")
add_code(c7, "X_train_sc = scaler.fit_transform(X_train)  # fit 仅训练集")
add_code(c7, "X_test_sc  = scaler.transform(X_test)        # 测试集只 transform")

# 3.2
add_sub(c7, '3.2 特征工程（三轮迭代）')
add_para(c7, '特征工程是本实验 AUC 从 0.944 提升至 0.950 的核心驱动，经历三轮迭代：', indent_first=True)

add_para(c7, '【第一轮】基础版本（AUC≈0.944）', bold=True)
add_para(c7, '仅完成 OHE + StandardScaler，使用原始58维特征直接训练。模型已具基础分类能力，但受限于特征的线性不可分性，AUC 停留在 0.944。', indent_first=True)

add_para(c7, '【第二轮】衍生特征工程（AUC≈0.946）', bold=True)
add_para(c7, 'previously_contacted（二值标志）：pdays=999 是数据集的特殊编码，表示"本次活动前从未联系过该客户"，包含重要的客户关系历史信息，将其提取为独立二值特征，并将 pdays=999 替换为 0（消除999这一大数值对标准化的干扰）。', indent_first=True)
add_code(c7, "df['previously_contacted'] = (df['pdays'] != 999).astype(int)")
add_code(c7, "df['pdays'] = df['pdays'].replace(999, 0)")
add_para(c7, 'log_duration（对数变换）：duration（通话时长）是最强单特征预测因子（均值约258秒，但最大值4918秒），分布呈强正偏态。取 log(duration+1) 可将偏态分布压缩为近似正态，增强逻辑回归的拟合质量。', indent_first=True)
add_code(c7, "df['log_duration'] = np.log1p(df['duration'])")

add_para(c7, '【第三轮】二阶交互特征（AUC≈0.9488~0.9503，最终方案）', bold=True)
add_para(c7, '逻辑回归是线性分类器，无法自动学习特征间的乘积交互效应。通过 PolynomialFeatures(degree=2, interaction_only=True) 为精选15列特征构造所有两两交互项（C(15,2)=105个），拼接到原始58维后共178维。', indent_first=True)
add_para(c7, '精选15列的依据：①duration 和 log_duration 是最强预测因子；②euribor3m、nr.employed、emp.var.rate、cons.price.idx、cons.conf.idx 是宏观经济状态的代理变量，与订阅决策高度相关且相互之间存在强交互；③poutcome_success 是已知的强正向信号；④previously_contacted 捕获客户历史关系。这些特征的两两乘积项可捕获如"在经济低迷期（euribor3m高）的长通话（duration大）更有可能转化"等交互效应。', indent_first=True)
add_code(c7, "top_feats = ['duration','log_duration','euribor3m','nr.employed',")
add_code(c7, "             'emp.var.rate','cons.price.idx','cons.conf.idx',")
add_code(c7, "             'previously_contacted','pdays','poutcome_success',")
add_code(c7, "             'poutcome_nonexistent','contact_cellular',")
add_code(c7, "             'month_may','month_nov','month_oct']")
add_code(c7, "poly = PolynomialFeatures(degree=2, interaction_only=True, include_bias=False)")
add_code(c7, "X_train_poly = poly.fit_transform(X_train_sc[:, idx])  # fit 仅训练集")
add_code(c7, "X_train_final = np.hstack([X_train_sc, X_train_poly])   # 原始+交互拼接")
add_para(c7, '配合 L1 正则化，模型自动从178个候选特征中稀疏选出有价值子集（实测约56~178个非零系数），既提升性能又控制过拟合。', indent_first=True)

# 3.3
add_sub(c7, '3.3 类别不平衡处理策略')
add_para(c7, '注意：所有重采样操作均仅在训练集上执行，测试集保持原始分布（11.27%），保证评估的公平性与现实意义。', indent_first=True)

add_para(c7, '策略A — SMOTE 过采样（Synthetic Minority Over-sampling Technique）', bold=True)
add_para(c7, '原理：对少数类中每个样本，在其 k 个近邻中随机插值合成新样本，而非简单重复。参数 sampling_strategy=0.4 控制最终正:负比约为2:5（约40933条），避免完全平衡时合成样本过多引入噪声，且训练集规模适中。', indent_first=True)
add_code(c7, "smote = SMOTE(sampling_strategy=0.4, random_state=42)")
add_code(c7, "X_tr_smote, y_tr_smote = smote.fit_resample(X_train_final, y_train)")
add_para(c7, '效果：Recall(Yes)=0.83，Precision(Yes)=0.54，Precision-Recall 间取得较好平衡，适合希望控制误报率的业务场景。', indent_first=True)

add_para(c7, '策略B — RandomUnderSampler 随机欠采样', bold=True)
add_para(c7, '原理：随机删除多数类样本，使正负样本数量相等（1:1，约7424条）。优点是训练集极小、速度快；缺点是丢弃大量真实负样本信息，约79%的原始训练数据被丢弃。', indent_first=True)
add_code(c7, "rus = RandomUnderSampler(random_state=42)")
add_code(c7, "X_tr_rus, y_tr_rus = rus.fit_resample(X_train_final, y_train)")
add_para(c7, '效果：Recall(Yes)=0.94（最高），Precision(Yes)=0.43，适合"宁可误报也不漏报"的高召回场景。', indent_first=True)

add_para(c7, '策略C — class_weight=\'balanced\'（损失权重调节）', bold=True)
add_para(c7, '原理：不修改数据，通过在损失函数中提高少数类权重来补偿不平衡。权重计算：weight_i = n_samples / (n_classes × n_samples_i)，即正样本的误分代价约为负样本的7.9倍。原始32950条训练数据完全保留，信息损失最小。', indent_first=True)
add_code(c7, "model = LogisticRegression(class_weight='balanced', C=0.1, penalty='l1',")
add_code(c7, "                           solver='liblinear', max_iter=3000)")
add_para(c7, '效果：Recall(Yes)=0.95（三者最高），AUC=0.9503（三者最高），是综合性能最优的策略。', indent_first=True)

# 3.4
add_sub(c7, '3.4 超参数调优（GridSearchCV + 5-Fold Stratified CV）')
add_para(c7, '参数网格设计：', indent_first=True)
add_code(c7, "param_grid = [")
add_code(c7, "  {'C':[0.005,0.01,0.02,0.05,0.1], 'penalty':['l1'], 'solver':['liblinear'], 'max_iter':[3000]},")
add_code(c7, "  {'C':[0.01,0.05,0.1,0.5,1],      'penalty':['l2'], 'solver':['lbfgs'],     'max_iter':[3000]},")
add_code(c7, "]   # 共10种参数组合 × 5折 = 50次拟合/每个策略")
add_para(c7, '核心参数解析：', indent_first=True)
add_para(c7, 'C（正则化强度倒数）：正则化项 = (1/C)·||w||。C越小，正则化越强，权重越趋近零，模型越简单。在178维高维特征空间中，过大的C会导致过拟合，最优C通常在0.05~1之间。', indent_first=True)
add_para(c7, 'penalty（正则化类型）：l1（Lasso）惩罚||w||₁，产生稀疏解，将无关特征权重压缩为0，等效自动特征选择，在高维场景下尤为关键；l2（Ridge）惩罚||w||₂，均匀压缩所有权重，适合特征高相关性场景。', indent_first=True)
add_para(c7, '为何用 roc_auc 而非 accuracy 作为 CV 评分：在正样本仅11.3%的场景下，将全部样本预测为负类可达88.7%的准确率，但此时的"高准确率"毫无意义。roc_auc 是阈值无关的指标，度量模型对正负样本的整体排序能力，不受类别不平衡影响。', indent_first=True)
add_code(c7, "cv = StratifiedKFold(n_splits=5, shuffle=True, random_state=42)")
add_code(c7, "gs = GridSearchCV(LogisticRegression(random_state=42, class_weight=cw),")
add_code(c7, "                  param_grid, cv=cv, scoring='roc_auc', n_jobs=-1, refit=True)")
add_code(c7, "gs.fit(X_tr, y_tr)")
add_code(c7, "best_model = gs.best_estimator_  # 自动在最佳参数下对全训练集重训练")

# ══════════════════════════════════════════════════════════════════════════════
# 四、实验结果及分析
# ══════════════════════════════════════════════════════════════════════════════
c8 = table.rows[8].cells[0]
cell_clear(c8)
add_para(c8, '四、实验结果及分析', bold=True)

# 4.1
add_sub(c8, '4.1 AUC 优化迭代对比')
add_para(c8, '本实验经过三轮迭代优化，AUC 从初始 0.944 提升至 0.950，共提升 0.006，以下表格记录了完整优化路径：', indent_first=True)

build_table(c8,
    headers=['迭代版本', '核心改进内容', '特征维度', 'SMOTE-AUC', 'RUS-AUC', 'CW-AUC'],
    rows=[
        ['迭代1（基础版）',    'OHE+StandardScaler，无衍生特征',                   '58',   '0.9436', '0.9438', '0.9442'],
        ['迭代2（衍生特征）',  '+previously_contacted, +log_duration, pdays修正', '60',   '0.9443', '0.9449', '0.9446'],
        ['迭代3（最终版）',    '+精选15列二阶交互项，L1自动特征选择',               '178',  '0.9499', '0.9488', '0.9503'],
        ['全量多项式（参考）', '全量58列二阶交互（仅CW，SMOTE内存溢出）',          '1711', '—',      '—',      '0.9528'],
    ],
    note='注：CW=class_weight=balanced；RUS=RandomUnderSampler；全量多项式因SMOTE在1711维×4万样本时内存溢出，仅CW策略可运行。'
)
add_para(c8, '')

# 4.2
add_sub(c8, '4.2 最终三模型完整性能对比（测试集，n=8238）')
build_table(c8,
    headers=['策略', 'Precision(Yes)', 'Recall(Yes)', 'F1(Yes)', 'Accuracy', 'AUC', 'TN/FP/FN/TP', '非零系数'],
    rows=[
        ['SMOTE',       '0.54', '0.83', '0.65', '90.0%', '0.9499', '6649/661/162/766',  '178'],
        ['UnderSample', '0.43', '0.94', '0.59', '85.0%', '0.9488', '6161/1149/58/870',   '56'],
        ['ClassWeight', '0.43', '0.95', '0.59', '85.0%', '0.9503', '6142/1168/50/878',  '108'],
    ],
    note='注：Yes类=已订阅正样本，测试集中support=928；非零系数数量反映L1正则的稀疏程度（178总特征中的有效特征数）。'
)
add_para(c8, '')

# 4.3
add_sub(c8, '4.3 评估指标深度分析')

add_para(c8, '（1）Accuracy 的局限性', bold=True)
add_para(c8, '若模型将全部测试样本预测为"未订阅"，准确率高达 88.7%（与班级同学的算法相当），但 Recall(Yes)=0，银行会错失全部4640位真实订阅客户。本实验三个模型的准确率为85%~90%，略低于"全预测负类"基准，但 Recall(Yes) 达到 83%~95%，才真正具有业务价值。', indent_first=True)

add_para(c8, '（2）三种策略的 Precision-Recall 权衡分析', bold=True)
add_para(c8, 'SMOTE（Recall=0.83，Precision=0.54）：在召回率和精确率间取得最佳平衡。误报661人、漏报162人，每预测10个"潜在订阅客户"中有5.4个是真实的，适合银行希望控制营销骚扰成本的场景。', indent_first=True)
add_para(c8, 'UnderSample（Recall=0.94，Precision=0.43）：召回极高，但误报1149人（占未订阅者15.7%）。大量负样本被删除后，模型过度倾向预测订阅，适合"不惜误报代价，确保找到所有潜在客户"的高召回需求。', indent_first=True)
add_para(c8, 'ClassWeight（Recall=0.95，Precision=0.43）：与欠采样结果相近，但保留全部32950条原始训练数据（信息损失最小），L1选出最有效的108个特征，且AUC在三者中最高（0.9503），是综合性能最优的策略。', indent_first=True)

add_para(c8, '（3）AUC 与最优参数分析', bold=True)
add_para(c8, 'GridSearchCV 最终选出的最优参数：SMOTE策略选 C=1, penalty=l2（数据经SMOTE平衡后，较强的L2正则更优）；UnderSample 和 ClassWeight 策略均选 C=0.05~0.1, penalty=l1（高维交互特征场景下，L1稀疏正则更适合进行特征选择）。三个模型 AUC 均达 0.9488~0.9503，远高于随机分类基准（0.5），验证了模型对潜在订阅客户具有显著的排序和识别能力。', indent_first=True)

add_para(c8, '（4）逻辑回归的性能上限与工程权衡', bold=True)
add_para(c8, '实验探索了全量二阶多项式特征（全部58维特征的两两交互，共1711维），在 class_weight 策略下AUC可达0.9528，但SMOTE在1711维×4万样本时产生内存溢出（约2.8GB，超出系统限制）。最终采用精选178维方案，在三种策略上均稳定运行，AUC达0.9488~0.9503，取得性能与内存效率的最优平衡。根据 Moro et al.（2014）原始论文，逻辑回归在此数据集上的理论上限约为0.95~0.96；要突破此上限须引入随机森林、XGBoost等可自动捕捉非线性交互的树模型。', indent_first=True)

# 4.4 混淆矩阵
add_sub(c8, '4.4 可视化分析')
add_para(c8, '图1  混淆矩阵对比（三种不平衡处理策略）', bold=False, align=WD_ALIGN_PARAGRAPH.CENTER)
if os.path.exists(IMG_CM):
    add_img(c8, IMG_CM, width_cm=15)
add_para(c8, '从混淆矩阵可直观看出三种策略的差异：SMOTE使FP（误报）最小（661），漏报（FN=162）相对较多；UnderSample 和 ClassWeight 的 FN 极低（58和50），但代价是 FP 显著上升（1149和1168），即为了不漏掉真正的订阅客户，会将更多无意愿客户也预测为"将订阅"。业务上，银行通常更担心漏报（错过真实客户），因此 ClassWeight 的低漏报策略（FN=50，漏报率5.4%）最具实践价值。', indent_first=True)

add_para(c8, '')
add_para(c8, '图2  ROC曲线对比（三种策略）', bold=False, align=WD_ALIGN_PARAGRAPH.CENTER)
if os.path.exists(IMG_ROC):
    add_img(c8, IMG_ROC, width_cm=12)
add_para(c8, '三条ROC曲线均高度贴近左上角（理想分类器），AUC差异仅0.0015（0.9488~0.9503），说明在整体排序能力上三种策略几乎一致，三条曲线均远高于对角线（随机分类器，AUC=0.5）基准。', indent_first=True)

# 4.5 结论
add_sub(c8, '4.5 实验结论')
conclusions = [
    '① 特征工程是提升线性模型性能的核心手段：通过三轮迭代（基础预处理→衍生特征→二阶交互特征），AUC从0.944稳步提升至0.950，验证了"特征质量决定模型上限"的机器学习实践原则；',
    '② 类别不平衡不可忽视：三种策略均能有效提升模型对少数类（订阅客户）的识别能力，Recall(Yes)从原始约0.40提升至0.83~0.95；其中 class_weight=balanced 在AUC（0.9503）和漏报率（5.4%）两项指标上均表现最优；',
    '③ L1 正则化在高维场景中具有双重价值：既控制过拟合，又通过稀疏化自动完成特征选择（178维中仅约56~108个系数非零），显著提升了模型的可解释性和泛化能力；',
    '④ 超参数调优发现最优 C 值较小（0.05~1），印证了高维特征空间下需要较强正则化的理论认知；调优评分指标应选 roc_auc 而非 accuracy，避免类别不平衡导致的评估误导；',
    '⑤ 逻辑回归在此数据集的性能上限约为AUC≈0.953（全量多项式特征+CW），若业务需要突破0.96，须引入能自动捕捉非线性交互的树模型（Random Forest、XGBoost等）。',
]
for c in conclusions:
    add_para(c8, c, indent_first=True)

# 保存
doc.save(OUT_PATH)
print(f"[OK] 报告已保存至: {OUT_PATH}")
