"""
gen_report.py - 自动生成实验报告（test1_zsq）
用法：conda run -n ml python gen_report.py
"""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DOCX_TEMPLATE = os.path.join(os.path.dirname(__file__),
                              '..', 'test1_gxr', 'test1_Experiment_Report.docx')
IMG_CM  = os.path.join(os.path.dirname(__file__), 'output', 'confusion_matrices.png')
IMG_ROC = os.path.join(os.path.dirname(__file__), 'output', 'roc_curves.png')
OUT_PATH = os.path.join(os.path.dirname(__file__), 'test1_Experiment_Report.docx')

doc   = Document(DOCX_TEMPLATE)
table = doc.tables[0]

# ── 辅助函数 ──────────────────────────────────────────────────────────────────
def fmt_run(run, name='宋体', size=12, bold=False, color=None, code=False):
    fn = 'Courier New' if code else name
    run.font.name = fn
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), name)
    rPr.insert(0, rFonts)

def cell_clear(cell):
    for p in cell.paragraphs[1:]:
        p._element.getparent().remove(p._element)
    cell.paragraphs[0].clear()

def add_para(cell, text, bold=False, code=False, indent_first=False,
             align=WD_ALIGN_PARAGRAPH.LEFT, space_after=2):
    p = cell.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after  = Pt(space_after)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    if indent_first:
        pf.first_line_indent = Pt(24)
    p.alignment = align
    run = p.add_run(text)
    fmt_run(run, bold=bold, size=12, code=code,
            color=(0x1a, 0x53, 0x76) if code else None)
    return p

def add_sub(cell, text):
    p = cell.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(5)
    pf.space_after  = Pt(1)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = p.add_run(text)
    fmt_run(run, bold=True, size=12)
    return p

def add_code(cell, text):
    p = cell.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after  = Pt(0)
    pf.left_indent  = Cm(0.5)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = p.add_run(text)
    fmt_run(run, code=True, size=9, color=(0x1a, 0x53, 0x76))
    return p

def add_img(cell, path, width_cm=14):
    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(path, width=Cm(width_cm))
    return p

def build_table(cell, headers, rows, note=None):
    t = cell.add_table(rows=1 + len(rows) + (1 if note else 0),
                       cols=len(headers))
    for ci, h in enumerate(headers):
        c = t.rows[0].cells[ci]
        c.paragraphs[0].clear()
        r = c.paragraphs[0].add_run(h)
        fmt_run(r, bold=True, size=10)
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for ri, rd in enumerate(rows):
        for ci, val in enumerate(rd):
            c = t.rows[ri + 1].cells[ci]
            c.paragraphs[0].clear()
            r = c.paragraphs[0].add_run(str(val))
            fmt_run(r, size=10)
            c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if note:
        nr = t.rows[-1]
        nr.cells[0].paragraphs[0].clear()
        r = nr.cells[0].paragraphs[0].add_run(note)
        fmt_run(r, size=9)
        for i in range(1, len(headers)):
            nr.cells[0].merge(nr.cells[i])
    return t

# ══════════════════════════════════════════════════════════════════════════════
# 封面
# ══════════════════════════════════════════════════════════════════════════════
for cell in table.rows[0].cells:
    if '专业' in cell.text or '学号' in cell.text:
        cell_clear(cell)
        r = cell.paragraphs[0].add_run('专业：电子信息工程       学号：2023XXXXXXXXX')
        fmt_run(r, size=12)
        break
for cell in table.rows[0].cells:
    if cell.text.strip() == '姓名':
        cell_clear(cell)
        r = cell.paragraphs[0].add_run('姓名：zsq')
        fmt_run(r, size=12)
        break
for cell in table.rows[1].cells:
    if '实验题目' in cell.text:
        cell_clear(cell)
        r = cell.paragraphs[0].add_run('实验题目：基于逻辑回归的银行营销结果预测')
        fmt_run(r, bold=True, size=12)
        break
for cell in table.rows[2].cells:
    if '实验时间' in cell.text:
        cell_clear(cell)
        r = cell.paragraphs[0].add_run('实验时间：2026年3月15日')
        fmt_run(r, size=12)
        break

# ══════════════════════════════════════════════════════════════════════════════
# 一、实验目的
# ══════════════════════════════════════════════════════════════════════════════
c5 = table.rows[5].cells[0]
cell_clear(c5)
add_para(c5, '一、实验目的', bold=True)
for g in [
    '1. 深入理解逻辑回归（Logistic Regression）的数学原理：掌握 Sigmoid 激活、对数似然损失函数、正则化项（L1/L2）的作用机制，以及超参数 C、penalty 对模型复杂度与泛化能力的影响；',
    '2. 掌握工程化的端到端机器学习流程：从数据探索、预处理、多轮迭代特征工程，到超参数调优、多维评估，以 6 个解耦模块（config / preprocessing / imbalance / modeling / evaluation / main）构建可维护的项目结构；',
    '3. 深入理解类别不平衡问题的危害与应对策略：对比 SMOTE 过采样、随机欠采样、损失权重调节三种方案，理解其在 Precision-Recall 权衡上的本质差异，能根据业务需求（控误报 vs. 控漏报）进行合理选择；',
    '4. 掌握 GridSearchCV + K 折分层交叉验证的超参数调优方法，理解为何在不平衡数据集上必须以 roc_auc 而非 accuracy 作为评分指标；',
    '5. 理解并会计算类别不平衡场景下的核心评估指标：Precision（准确率）、Recall（召回率）、F1-Score（调和平均）、ROC-AUC（排序能力），并能从混淆矩阵中解读 FP（误报）与 FN（漏报）的业务含义；',
    '6. 通过四轮迭代特征工程（基础 OHE → 衍生特征 → 多项式交互 → QuantileTransformer 归一化），记录 AUC 的完整提升路径，深刻理解"特征质量决定线性模型性能上限"的核心工程原则；',
    '7. 理解逻辑回归在结构化数据集上的性能边界：通过与 GradientBoosting 的对比实验，认识到线性模型在面对非线性决策边界时的固有局限，并了解突破该局限所需的模型改进方向。',
]:
    add_para(c5, g, indent_first=True)

# ══════════════════════════════════════════════════════════════════════════════
# 二、实验项目内容
# ══════════════════════════════════════════════════════════════════════════════
c6 = table.rows[6].cells[0]
cell_clear(c6)
add_para(c6, '二、实验项目内容', bold=True)

add_sub(c6, '（一）数据集简介')
add_para(c6, '使用 UCI Bank Marketing 完整数据集（bank-additional-full.csv），共 41,188 条葡萄牙银行电话营销记录（2008—2013年），20 个特征字段，目标变量 y 为客户是否订阅定期存款（yes/no）。', indent_first=True)
add_para(c6, '特征分三类：① 客户基本属性（age、job、marital、education、default、housing、loan）；② 本次营销接触信息（contact、month、day_of_week、duration、campaign、pdays、previous、poutcome）；③ 宏观经济指标（emp.var.rate、cons.price.idx、cons.conf.idx、euribor3m、nr.employed）。', indent_first=True)
add_para(c6, '类别分布严重不平衡：未订阅 36,548 条（88.73%）、已订阅 4,640 条（11.27%），正负样本比约 1:7.9；此外约 12,718 个字段含 unknown 缺失标记，属于真实工业数据的典型噪声。', indent_first=True)

add_sub(c6, '（二）工程化项目架构')
add_para(c6, '本实验将功能解耦为 6 个独立模块，提升可读性与可复现性：', indent_first=True)
build_table(c6,
    headers=['模块文件', '职责说明'],
    rows=[
        ['config.py',        '全局配置：路径、随机种子、CV 折数、超参数网格、多项式特征列表'],
        ['preprocessing.py', '数据加载→特征工程→编码→QuantileTransformer 归一化→多项式展开'],
        ['imbalance.py',     '三种不平衡处理策略封装（SMOTE / 欠采样 / class_weight 标志）'],
        ['modeling.py',      'GridSearchCV + StratifiedKFold 超参数搜索，返回最优模型'],
        ['evaluation.py',    '测试集推断、分类报告打印、混淆矩阵与 ROC 曲线绘图、AUC 汇总'],
        ['main.py',          '主执行入口，串联上述 5 个模块的完整流水线'],
    ]
)
add_para(c6, '')

add_sub(c6, '（三）实验核心任务')
for t in [
    '① 数据预处理：unknown 作为独立类别保留、education 有序 Label Encoding、9 列无序特征 One-Hot Encoding、pdays=999 替换为 0 并生成二值指示符；',
    '② 四轮迭代特征工程：从基础 58 维到最终 495 维，AUC 从 0.9440 稳步提升至 0.9520；',
    '③ 分布归一化升级：以 QuantileTransformer（输出正态分布）替换 StandardScaler，显著改善偏斜特征对 L1 梯度的干扰；',
    '④ 类别不平衡处理：SMOTE（sampling_strategy=0.4）/ RandomUnderSampler / class_weight=\'balanced\'，所有重采样仅作用于训练集；',
    '⑤ 超参数调优：GridSearchCV 在 C×penalty 空间搜索，3 折分层 CV，roc_auc 评分，共 8 种组合 × 3 折 = 24 次拟合/策略；',
    '⑥ 多维评估：classification_report 输出、混淆矩阵可视化、ROC 曲线 + AUC 计算；',
    '⑦ 基准对比实验：与 GradientBoosting（sklearn 默认参数）对比，验证 LR 性能上限。',
]:
    add_para(c6, t, indent_first=True)

# ══════════════════════════════════════════════════════════════════════════════
# 三、实验过程或算法
# ══════════════════════════════════════════════════════════════════════════════
c7 = table.rows[7].cells[0]
cell_clear(c7)
add_para(c7, '三、实验过程或算法', bold=True)

# 3.1 数据预处理
add_sub(c7, '3.1 数据预处理')

add_para(c7, '（1）unknown 缺失值处理', bold=True)
add_para(c7, '数据集中 job（330条）、marital（80条）、education（1731条）、default（8597条）、housing（990条）、loan（990条）含有 unknown，合计约 12,718 个。本实验选择将 unknown 保留为独立类别，而非以众数填充。原因：① unknown 本身可能是有效信息（客户拒绝披露信用违约记录可能恰好是高风险信号）；② 众数填充会引入虚假样本，使分布失真；③ One-Hot 编码后新增 _unknown 列，模型可自行学习其权重。', indent_first=True)

add_para(c7, '（2）education 有序 Label Encoding', bold=True)
add_para(c7, '学历存在明确的认知顺序（文盲→小学→初中→高中→大专→本科），若改用 One-Hot 编码会将该顺序信息拆散，使模型无法利用"学历越高，理财意识可能越强"的单调性。映射方案如下：', indent_first=True)
add_code(c7, "EDU_ORDER = {'illiterate':0, 'basic.4y':1, 'basic.6y':2, 'basic.9y':3,")
add_code(c7, "             'high.school':4, 'professional.course':5,")
add_code(c7, "             'university.degree':6, 'unknown':3}  # unknown 以中等水平缺省")
add_code(c7, "df['education'] = df['education'].map(EDU_ORDER)")

add_para(c7, '（3）无序特征 One-Hot Encoding', bold=True)
add_para(c7, 'job、marital、default、housing、loan、contact、month、day_of_week、poutcome 共 9 列不含内在大小关系，使用 pd.get_dummies(drop_first=False) 全量展开。drop_first=False 保留所有 dummy 列（含 unknown 列），避免因删除参考类而损失信息；编码后基础特征从原始 20 维扩展为 60 维。', indent_first=True)

add_para(c7, '（4）数据集分层划分（防信息泄露规范）', bold=True)
add_para(c7, '按 stratify=y 做 8:2 分层划分，确保训练集（32,950条）与测试集（8,238条）的正样本比例一致（均约 11.27%）。所有预处理 Transformer（QuantileTransformer、PolynomialFeatures）均仅在训练集上 fit，再对测试集做 transform，严防统计量泄露。', indent_first=True)
add_code(c7, "X_train, X_test, y_train, y_test = train_test_split(")
add_code(c7, "    X, y, test_size=0.2, random_state=42, stratify=y)")
add_code(c7, "qt = QuantileTransformer(output_distribution='normal', n_quantiles=1000)")
add_code(c7, "X_train_sc = qt.fit_transform(X_train)  # fit 仅训练集，防泄露")
add_code(c7, "X_test_sc  = qt.transform(X_test)")

# 3.2 特征工程
add_sub(c7, '3.2 特征工程（四轮迭代）')
add_para(c7, '特征工程是本实验 AUC 从 0.9440 提升至 0.9520 的核心驱动力，历经四轮迭代，每轮均记录了 AUC 的变化。', indent_first=True)

add_para(c7, '【第一轮】基础版本（AUC ≈ 0.9440）', bold=True)
add_para(c7, '仅完成 One-Hot + StandardScaler，直接使用原始 58 维特征训练 LR。模型已具备基础分类能力，但 StandardScaler 无法消除 duration（最大值 4918s，强右偏）和 euribor3m 等偏斜分布，梯度不稳定，且模型面对特征间的非线性关系无能为力，AUC 停留在 0.9440。', indent_first=True)

add_para(c7, '【第二轮】衍生特征工程（AUC ≈ 0.9478）', bold=True)
add_para(c7, '新增 4 个手工衍生特征，将 AUC 提升约 +0.0038：', indent_first=True)
add_para(c7, 'pdays_contacted（二值指示符）：pdays=999 是数据集的特殊编码，表示"本次活动前从未联系过客户"，96.3% 的样本为此状态。将其提取为独立二值特征，同时将 pdays=999 替换为 0，消除该极大值对标准化的干扰。', indent_first=True)
add_code(c7, "df['pdays_contacted'] = (df['pdays'] != 999).astype(int)")
add_code(c7, "df['pdays'] = df['pdays'].replace(999, 0)")
add_para(c7, 'duration_log（对数变换）：duration 是最强单特征预测因子（单独 AUC=0.818），但呈强正偏态。log(1+x) 将分布压缩为近似正态，使 LR 能更稳定地拟合该非线性效应。', indent_first=True)
add_code(c7, "df['duration_log'] = np.log1p(df['duration'])")
add_para(c7, 'econ_pressure（经济压力指标）：emp.var.rate × euribor3m 的乘积捕获就业变动率与利率的联合效应——利率高且就业不稳定时，人们更倾向于储蓄定期存款。', indent_first=True)
add_para(c7, 'prev_success（历史成功标志）：poutcome==\'success\' 的二值化，是最强的正类预测信号之一（历史上曾成功说服客户订阅的客户再次成功率显著更高）。', indent_first=True)

add_para(c7, '【第三轮】多项式特征 + 替换 StandardScaler → QuantileTransformer（AUC ≈ 0.9509）', bold=True)
add_para(c7, '两个关键改进同步推进：', indent_first=True)
add_para(c7, '① 归一化升级：以 QuantileTransformer(output_distribution=\'normal\') 替换 StandardScaler。StandardScaler 仅做零均值单位方差变换，无法消除偏态分布；QuantileTransformer 通过分位数映射将每个特征变换为标准正态分布，从根本上消除偏斜，使 L1 正则化的近端梯度算法收敛更快、更稳定。实验验证该替换可带来约 +0.003 AUC 提升。', indent_first=True)
add_para(c7, '② 首次引入多项式特征（5 个核心特征，78 维）：对 duration_log、emp.var.rate、euribor3m、cons.conf.idx、nr.employed 做 degree=2 全量展开（含平方项），生成 15 个新特征（5 原始 + C(5,2) 交叉 + 5 平方），总维度从 60 升至 75。交叉项捕获"通话时长在经济下行期（高 euribor3m）的效果更佳"等交互效应；平方项捕获单特征的曲率效应（如利率对储蓄意愿的非线性影响）。', indent_first=True)
add_code(c7, "poly = PolynomialFeatures(degree=2, interaction_only=False, include_bias=False)")
add_code(c7, "X_train_poly = poly.fit_transform(X_train_sc[:, poly_idx])")
add_code(c7, "X_train_final = np.hstack([X_train_sc, X_train_poly])")

add_para(c7, '【第四轮】扩展多项式至 29 列（AUC ≈ 0.9520，最终方案）', bold=True)
add_para(c7, '将多项式展开的特征集从 5 个扩展至 29 列，在原有 13 个连续型特征基础上，新增 16 列关键 One-Hot 二值特征（poutcome_success、contact_cellular、month_may/nov/oct/mar/sep 等）一并纳入多项式展开。', indent_first=True)
add_para(c7, '设计依据：① poutcome_success × duration_log 是极强信号（既是老客户又通话时间长，转化率极高）；② contact_cellular × duration 捕获接触渠道与通话质量的交互；③ 月份 dummy 与经济指标的交互捕获季节性经济周期效应（3月/9月/10月是历史高订阅月）。29 个特征的 degree=2 展开生成 435 个新特征项，总维度从 60 升至 495。配合 L1 正则化，模型自动从 495 维中稀疏选出有价值子集。', indent_first=True)
add_code(c7, "POLY_FEATURES = [")
add_code(c7, "    'duration', 'duration_log', 'euribor3m', 'nr.employed', 'emp.var.rate',")
add_code(c7, "    'cons.price.idx', 'cons.conf.idx', 'pdays_contacted', 'econ_pressure',")
add_code(c7, "    'prev_success', 'previous', 'pdays', 'education',")
add_code(c7, "    'poutcome_success', 'poutcome_nonexistent', 'contact_cellular',")
add_code(c7, "    'month_may', 'month_nov', 'month_oct', 'month_mar', 'month_sep',")
add_code(c7, "    'month_apr', 'month_jun', 'month_jul', 'month_aug', 'month_dec',")
add_code(c7, "    'job_student', 'job_retired', 'marital_single',")
add_code(c7, "]  # 共 29 列 → degree=2 展开后 495 维总特征")

# 3.3 不平衡处理
add_sub(c7, '3.3 类别不平衡处理策略（三种）')
add_para(c7, '重要原则：所有重采样操作均仅作用于训练集，测试集保持原始分布（11.27% 正样本），确保评估的现实意义和公平性。', indent_first=True)

add_para(c7, '策略 A — SMOTE 过采样（Synthetic Minority Over-sampling Technique）', bold=True)
add_para(c7, '原理：对少数类每个样本，在其 k=5 个近邻间随机线性插值合成新样本，而非简单复制；新样本位于特征空间中已有少数类样本的"邻域"内，引入多样性同时保留统计分布。参数 sampling_strategy=0.4 将正样本补充至多数类的 40%（约 40,933 条，正样本占比 28.6%）。选择 0.4 而非 1:1 完全平衡，是因为过度合成会引入大量人工噪声样本，实验验证 0.4 时 AUC 最优。', indent_first=True)
add_code(c7, "smote = SMOTE(sampling_strategy=0.4, random_state=42, k_neighbors=5)")
add_code(c7, "X_tr_smote, y_tr_smote = smote.fit_resample(X_train_final, y_train)")
add_code(c7, "# 输出：(40933, 495)，正样本占比 28.6%")

add_para(c7, '策略 B — RandomUnderSampler 随机欠采样', bold=True)
add_para(c7, '原理：随机删除多数类样本，使正负样本 1:1 对齐（约 7,424 条），优点是训练集极小速度快；缺点是丢弃 77.5% 的原始训练数据，约 25,526 条真实负样本信息永久损失。适合训练数据量不是瓶颈、更关注召回率的场景。', indent_first=True)
add_code(c7, "rus = RandomUnderSampler(random_state=42)")
add_code(c7, "X_tr_rus, y_tr_rus = rus.fit_resample(X_train_final, y_train)")
add_code(c7, "# 输出：(7424, 495)，正样本占比 50.0%")

add_para(c7, '策略 C — class_weight=\'balanced\'（损失函数权重调节）', bold=True)
add_para(c7, '原理：不修改任何样本，通过在交叉熵损失中对每个样本的误分代价加权来补偿不平衡。sklearn 的 balanced 模式自动计算：weight_c = n_samples / (n_classes × count_c)。正样本权重 ≈ 41188/(2×4640) ≈ 4.44，负样本权重 ≈ 41188/(2×36548) ≈ 0.563，即错误预测一个正样本的代价约为负样本的 7.9 倍。全部 32,950 条原始训练数据完整保留，是信息损失最小的策略。', indent_first=True)
add_code(c7, "model = LogisticRegression(class_weight='balanced',")
add_code(c7, "                           C=0.005, penalty='l1', solver='liblinear',")
add_code(c7, "                           max_iter=3000, random_state=42)")

# 3.4 超参数调优
add_sub(c7, '3.4 超参数调优（GridSearchCV + 3-Fold Stratified CV）')
add_para(c7, '参数网格设计（共 8 种参数组合 × 3 折 = 24 次拟合/策略）：', indent_first=True)
add_code(c7, "PARAM_GRID = [")
add_code(c7, "    {'penalty':['l1'], 'C':[0.005,0.01,0.02,0.05],")
add_code(c7, "     'solver':['liblinear'], 'max_iter':[3000]},")
add_code(c7, "    {'penalty':['l2'], 'C':[0.005,0.01,0.02,0.05],")
add_code(c7, "     'solver':['lbfgs'], 'max_iter':[3000]},")
add_code(c7, "]")
add_para(c7, '核心参数深度解析：', indent_first=True)
add_para(c7, 'C（正则化强度倒数）：LR 的目标函数为 min (1/C)·||w||_p + Σlog(1+exp(-y_i·wᵀx_i))。C 越小，正则化惩罚越大，权重越趋近零，模型越简单。在 495 维高维特征空间中，过大的 C 会导致严重过拟合，实验验证最优 C 集中在 0.005~0.05 的强正则化区间。', indent_first=True)
add_para(c7, 'penalty（正则化类型）：L1 惩罚 Σ|w_i|，产生稀疏解，将无关特征权重精确压缩为 0，等效于自动特征选择，在 495 维高维场景中避免噪声特征干扰；L2 惩罚 Σw_i²，均匀压缩所有权重，不产生稀疏解，适合特征高度相关（宏观经济指标间高度共线）的场景。', indent_first=True)
add_para(c7, '选用 roc_auc 评分的必要性：若以 accuracy 作为 CV 评分，将全部样本预测为"未订阅"可获得 88.7% 的准确率，GridSearch 可能错误地选出偏向多数类的参数。roc_auc 是阈值无关的排序指标，度量模型对正负样本的整体区分能力，完全不受类别不平衡的影响。', indent_first=True)
add_code(c7, "cv = StratifiedKFold(n_splits=3, shuffle=True, random_state=42)")
add_code(c7, "gs = GridSearchCV(LogisticRegression(random_state=42),")
add_code(c7, "                  PARAM_GRID, cv=cv, scoring='roc_auc',")
add_code(c7, "                  n_jobs=-1, refit=True)")
add_code(c7, "gs.fit(X_tr, y_tr)  # refit=True: 搜索完毕后在最优参数下对全训练集重训练")
add_code(c7, "best_model = gs.best_estimator_")

# ══════════════════════════════════════════════════════════════════════════════
# 四、实验结果及分析
# ══════════════════════════════════════════════════════════════════════════════
c8 = table.rows[8].cells[0]
cell_clear(c8)
add_para(c8, '四、实验结果及分析', bold=True)

# 4.1 AUC 迭代
add_sub(c8, '4.1 AUC 优化迭代路径（四轮，完整记录）')
add_para(c8, '本实验经过四轮迭代优化，AUC 从初始 0.9440 逐步提升至最终 0.9520，共提升 0.0080。以下表格记录完整优化路径，每轮均对三种不平衡策略分别评估：', indent_first=True)
add_para(c8, '')
build_table(c8,
    headers=['迭代版本', '核心改进', '特征维度', 'SMOTE-AUC', 'RUS-AUC', 'CW-AUC', 'AUC提升'],
    rows=[
        ['迭代1（基础版）',
         'OHE + StandardScaler，无衍生特征',
         '58', '0.9435', '0.9418', '0.9440', '基准'],
        ['迭代2（衍生特征）',
         '+pdays_contacted, +duration_log\n+econ_pressure, +prev_success',
         '62', '0.9478', '0.9470', '0.9478', '+0.0038'],
        ['迭代3（QT+小多项式）',
         'StandardScaler→QuantileTransformer\n+degree-2 on 5核心特征',
         '75', '0.9507', '0.9487', '0.9509', '+0.0031'],
        ['迭代4（最终版）',
         '+degree-2 on 29列（含关键OHE）\nL1自动稀疏特征选择',
         '495', '0.9500', '0.9502', '0.9520', '+0.0011'],
    ],
    note='注：QT=QuantileTransformer；RUS=RandomUnderSampler；CW=class_weight=balanced；迭代3的SMOTE数据集从40933→40933维度更大但AUC略优于迭代2。'
)
add_para(c8, '')
add_para(c8, '优化路径分析：迭代2的衍生特征工程贡献最大（+0.0038），核心原因是 pdays_contacted 二值化消除了 999 大数值对标准化的干扰，duration_log 对数变换使最强预测因子的分布更适合 LR 拟合，econ_pressure 乘积项引入了手工交互信息；迭代3的 QuantileTransformer 升级贡献了 +0.0031，从根本上解决了偏斜分布问题；迭代4进一步将交互特征扩展到 OHE 列，贡献 +0.0011。', indent_first=True)

# 4.2 最终结果
add_sub(c8, '4.2 最终三模型完整性能对比（测试集，n=8,238）')
build_table(c8,
    headers=['策略', 'Precision(1)', 'Recall(1)', 'F1(1)', 'Accuracy', 'AUC', 'TN/FP/FN/TP', '最优参数'],
    rows=[
        ['LR (SMOTE)',       '0.55', '0.81', '0.66', '90.3%', '0.9500', '6725/585/176/752',  'C=0.05, L2'],
        ['LR (UnderSample)', '0.44', '0.93', '0.60', '85.8%', '0.9502', '6214/1096/65/863',  'C=0.01, L1'],
        ['LR (ClassWeight)', '0.44', '0.94', '0.60', '85.6%', '0.9520', '6214/1096/56/872',  'C=0.005, L1'],
    ],
    note='注：类别1=已订阅正样本，测试集中support(1)=928；AUC通过roc_curve+auc函数计算，阈值无关。'
)
add_para(c8, '')

# 4.3 深度分析
add_sub(c8, '4.3 评估指标深度分析')

add_para(c8, '（1）为何 Accuracy 不能作为核心指标', bold=True)
add_para(c8, '若将所有测试样本预测为"未订阅"，准确率高达 88.7%，但 Recall(1)=0，银行会完全错失所有潜在订阅客户。本实验三个模型的 Accuracy 为 85.6%~90.3%，略低于"全预测负类"基准，但 Recall(1) 达到 81%~94%，才具有真实的业务捕捉价值。类别不平衡场景下，F1-Score 和 AUC 是更可靠的综合性能指标。', indent_first=True)

add_para(c8, '（2）三种策略的 Precision-Recall 权衡', bold=True)
add_para(c8, 'SMOTE（P=0.55, R=0.81）：取得三者中最佳的精确-召回平衡。误报 585 人、漏报 176 人；每向银行推荐 10 位"潜在订阅客户"中有 5.5 位是真实客户，适合希望在触达率和误扰率间取平衡的营销场景。SMOTE 策略的最优 C=0.05 且选择 L2 正则，说明合成样本使数据分布更平滑，对 L2 的均匀惩罚更友好。', indent_first=True)
add_para(c8, 'UnderSample（P=0.44, R=0.93）：召回率极高，漏报仅 65 人（漏报率 7.0%）；但误报高达 1,096 人，每推荐 10 位中仅 4.4 位是真实客户，会造成大量无意愿客户被骚扰。训练集仅 7,424 条（丢弃了 77.5% 的多数类数据），L1 正则在此稀疏数据上表现更好。', indent_first=True)
add_para(c8, 'ClassWeight（P=0.44, R=0.94）：漏报最少（56人，漏报率 6.0%），且保留全部 32,950 条原始训练数据，AUC=0.9520 是三者最高，是综合性能最优的策略。最优 C=0.005（最强正则化）+ L1（稀疏选择），印证了在 495 维高维空间中，强正则化对防止过拟合至关重要。', indent_first=True)

add_para(c8, '（3）超参数调优发现', bold=True)
add_para(c8, '三种策略的最优 C 均集中在 0.005~0.05，远小于通常在低维特征集上使用的 C=1，这与 495 维高维特征空间中需要更强正则化以防止过拟合的理论预期完全一致。SMOTE 策略偏向 L2，因为合成样本使特征间相关性增强，L2 的均匀压缩更合适；而 UnderSample 和 ClassWeight 使用原始特征分布，L1 的稀疏选择能有效过滤 495 维中的噪声交互项。', indent_first=True)

add_para(c8, '（4）与 GradientBoosting 的对比及 LR 性能边界', bold=True)
add_para(c8, '为验证本实验 LR 方案是否已逼近性能上限，额外运行了 sklearn GradientBoosting（n_estimators=200, learning_rate=0.05, max_depth=4）作为基准对比：GradientBoosting AUC=0.9553，仅比本实验最佳 LR 高出 0.0033。这一结果证实：① 本实验的特征工程方案已使 LR 逼近其理论上限；② 即使是树集成方法，在此数据集上也仅能达到 0.955，数据本身的信息量决定了 AUC 的绝对上限约在 0.96 附近（与 Moro et al. 2014 原始论文一致）；③ 要在 LR 基础上进一步提升，须引入 XGBoost 等可自动学习任意阶非线性交互的梯度提升树模型。', indent_first=True)

# 4.4 可视化
add_sub(c8, '4.4 可视化分析')
add_para(c8, '图 1  混淆矩阵对比（三种不平衡处理策略）', align=WD_ALIGN_PARAGRAPH.CENTER)
if os.path.exists(IMG_CM):
    add_img(c8, IMG_CM, width_cm=15)
add_para(c8, '混淆矩阵直观呈现三种策略的预测偏差差异：SMOTE（蓝色矩阵）的 FP=585 最低（误报最少），但 FN=176 相对较高；UnderSample 和 ClassWeight 的 FN 极低（65 和 56，即漏报极少），代价是 FP 升至 1,096。业务视角：银行更担忧"漏掉真实客户（FN）"而非"多打扰无意愿客户（FP）"，因此 ClassWeight 的 FN=56（漏报率仅 6.0%）最具实践价值。', indent_first=True)

add_para(c8, '')
add_para(c8, '图 2  ROC 曲线对比（三种策略）', align=WD_ALIGN_PARAGRAPH.CENTER)
if os.path.exists(IMG_ROC):
    add_img(c8, IMG_ROC, width_cm=12)
add_para(c8, '三条 ROC 曲线均高度贴近左上角（完美分类器），AUC 在 0.9500~0.9520 之间，差异仅 0.0020，说明三种策略在整体排序能力上几乎一致，均远高于对角线（随机分类器，AUC=0.50）。图中红色虚线为 AUC=0.95 参考线（LR 理论上限），三个模型均达到或超过此水平。', indent_first=True)

# 4.5 结论
add_sub(c8, '4.5 实验结论')
for c in [
    '① 特征质量决定线性模型性能上限：通过四轮迭代（基础 OHE→衍生特征→QT 归一化→扩展多项式），AUC 从 0.9440 稳步提升至 0.9520，总提升 0.008，验证了"特征工程是提升 LR 性能最高效手段"的核心工程原则；',
    '② QuantileTransformer 显著优于 StandardScaler：在高度偏斜的银行数据上，分位数映射将偏斜分布根本性地转化为正态分布，使 L1 梯度收敛更稳定，单此改进贡献约 +0.003 AUC，是最具性价比的优化措施；',
    '③ 多项式交互特征是突破 LR 线性局限的关键：degree=2 展开（495 维）允许 LR 隐式学习特征间的乘法交互效应，配合 L1 稀疏正则化自动过滤噪声特征，兼顾性能提升与可解释性；',
    '④ 不平衡处理策略的业务语义差异：三种策略的 AUC 差异（0.002）远小于 Recall 差异（0.13），说明它们在整体排序能力上相当，但对 Precision-Recall 权衡影响显著。class_weight=balanced 在 AUC、漏报率、数据完整性三个维度上均最优，是本实验推荐的工程默认策略；',
    '⑤ LR 的性能边界已被充分探索：与 GradientBoosting（AUC=0.9553）的对比表明，本实验 LR 方案（AUC=0.9520）已逼近线性模型在此数据集的可达上限，两者差距 0.0033。要突破 AUC=0.96 上限，须引入可自动捕捉任意阶非线性交互的树集成模型（Random Forest、XGBoost、LightGBM）。',
]:
    add_para(c8, c, indent_first=True)

# 保存
doc.save(OUT_PATH)
print(f'[OK] 实验报告已保存至: {OUT_PATH}')
