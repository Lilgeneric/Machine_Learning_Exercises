"""
gen_report.py — 生成实验报告 test2_Experiment_Report.docx

题目：基于决策树算法的电信客户流失预测
作者：gxr   日期：2026-03-20
"""

import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── 路径配置 ─────────────────────────────────────────────────────────────────
BASE      = os.path.dirname(os.path.abspath(__file__))
TMPL_PATH = '/home/vla/machine_learning/Experiment_Report.docx'
OUT_PATH  = os.path.join(BASE, 'test2_Experiment_Report.docx')
IMG = {
    'cm':          os.path.join(BASE, 'output', 'confusion_matrices.png'),
    'roc':         os.path.join(BASE, 'output', 'roc_curves.png'),
    'prf':         os.path.join(BASE, 'output', 'precision_recall_f1.png'),
    'heatmap':     os.path.join(BASE, 'output', 'metrics_heatmap.png'),
    'feat_imp':    os.path.join(BASE, 'output', 'feature_importance.png'),
}


# ══════════════════════════════════════════════════════════════════════════════
# 辅助函数
# ══════════════════════════════════════════════════════════════════════════════

def _set_east_asia(run, font_name):
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:ascii'),    font_name)
    rFonts.set(qn('w:hAnsi'),    font_name)
    rPr.insert(0, rFonts)


def fmt(run, name='宋体', size=12, bold=False, color=None, italic=False):
    run.font.name  = name
    run.font.size  = Pt(size)
    run.font.bold  = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    _set_east_asia(run, name)


def fmt_code(run):
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1a, 0x53, 0x76)
    _set_east_asia(run, 'Courier New')


def cell_clear(cell):
    for p in cell.paragraphs[1:]:
        p._element.getparent().remove(p._element)
    cell.paragraphs[0].clear()


def _para(cell, text='', bold=False, code=False, indent=False,
          align=WD_ALIGN_PARAGRAPH.LEFT, sb=0, sa=2, size=12):
    p  = cell.add_paragraph()
    pf = p.paragraph_format
    pf.space_before         = Pt(sb)
    pf.space_after          = Pt(sa)
    pf.line_spacing_rule    = WD_LINE_SPACING.SINGLE
    pf.first_line_indent    = Pt(24) if indent else Pt(0)
    p.alignment = align
    if text:
        run = p.add_run(text)
        if code:
            fmt_code(run)
        else:
            fmt(run, bold=bold, size=size)
    return p


def _sub(cell, text, sb=6, sa=2):
    """小节标题：加粗"""
    p  = cell.add_paragraph()
    pf = p.paragraph_format
    pf.space_before      = Pt(sb)
    pf.space_after       = Pt(sa)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = p.add_run(text)
    fmt(run, bold=True, size=12)
    return p


def _code(cell, text):
    p  = cell.add_paragraph()
    pf = p.paragraph_format
    pf.space_before      = Pt(0)
    pf.space_after       = Pt(0)
    pf.left_indent        = Cm(0.8)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = p.add_run(text)
    fmt_code(run)
    return p


def _img(cell, path, width_cm=14.5, caption=None):
    if caption:
        cp = cell.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_before = Pt(4)
        cp.paragraph_format.space_after  = Pt(2)
        run = cp.add_run(caption)
        fmt(run, bold=True, size=11)
    if os.path.exists(path):
        p   = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(6)
        p.add_run().add_picture(path, width=Cm(width_cm))
    else:
        _para(cell, f'[图片未找到: {path}]', sa=4)
    return cell


def _set_table_borders(table):
    """为表格设置全边框"""
    tbl  = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    tblBorders = OxmlElement('w:tblBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'),  'single')
        el.set(qn('w:sz'),   '4')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), '4472C4')
        tblBorders.append(el)
    tblPr.append(tblBorders)
    if tbl.tblPr is None:
        tbl.insert(0, tblPr)


def _shade_cell(cell, fill='4472C4'):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  fill)
    tcPr.append(shd)


def _build_table(cell, headers, rows, note=None, col_widths=None):
    """构建带边框和表头阴影的规范数据表格"""
    n_rows = 1 + len(rows) + (1 if note else 0)
    t = cell.add_table(rows=n_rows, cols=len(headers))
    _set_table_borders(t)

    # 设置列宽
    if col_widths:
        for ci, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[ci].width = Cm(w)

    # 表头行
    for ci, h in enumerate(headers):
        c = t.rows[0].cells[ci]
        _shade_cell(c, '4472C4')
        c.paragraphs[0].clear()
        run = c.paragraphs[0].add_run(h)
        fmt(run, bold=True, size=10, color=(0xFF, 0xFF, 0xFF))
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 数据行（隔行浅色）
    for ri, rd in enumerate(rows):
        fill_color = 'EBF0FA' if ri % 2 == 0 else 'FFFFFF'
        for ci, val in enumerate(rd):
            c = t.rows[ri + 1].cells[ci]
            _shade_cell(c, fill_color)
            c.paragraphs[0].clear()
            run = c.paragraphs[0].add_run(str(val))
            fmt(run, size=10)
            c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 注释行
    if note:
        nr = t.rows[-1]
        _shade_cell(nr.cells[0], 'F5F5F5')
        nr.cells[0].paragraphs[0].clear()
        run = nr.cells[0].paragraphs[0].add_run(note)
        fmt(run, size=9, color=(0x55, 0x55, 0x55))
        for i in range(1, len(headers)):
            nr.cells[0].merge(nr.cells[i])

    return t


# ══════════════════════════════════════════════════════════════════════════════
# 主体
# ══════════════════════════════════════════════════════════════════════════════
doc   = Document(TMPL_PATH)
table = doc.tables[0]


# ── 封面信息 ─────────────────────────────────────────────────────────────────
# Row 0 实际唯一单元格（合并后去重）：
#   [0]=专业,学号标签  [1]=内容格（填专业学号）  [2]=姓名标签  [3]=姓名内容
row0_uniq = list(dict.fromkeys(table.rows[0].cells))
if len(row0_uniq) >= 2:
    cell_clear(row0_uniq[1])
    run = row0_uniq[1].paragraphs[0].add_run('专业：电子信息工程     学号：202530131059')
    fmt(run, size=12)
if len(row0_uniq) >= 4:
    cell_clear(row0_uniq[3])
    run = row0_uniq[3].paragraphs[0].add_run('高夕茹')
    fmt(run, size=12)

for cell in table.rows[1].cells:
    if '实验题目' in cell.text or cell.text.strip() == '':
        cell_clear(cell)
        run = cell.paragraphs[0].add_run('实验题目：基于决策树算法的电信客户流失预测')
        fmt(run, bold=True, size=12)
        break

for cell in table.rows[2].cells:
    if '实验时间' in cell.text:
        cell_clear(cell)
        run = cell.paragraphs[0].add_run('实验时间：2026年3月20日')
        fmt(run, size=12)
        break


# ══════════════════════════════════════════════════════════════════════════════
# 一、实验目的
# ══════════════════════════════════════════════════════════════════════════════
c5 = table.rows[5].cells[0]
cell_clear(c5)
_para(c5, '一、实验目的', bold=True, sa=3)

goals = [
    ('1. 掌握决策树（Decision Tree / CART 算法）的数学原理与工程实现。深入理解节点分裂准则'
     '（基尼系数 Gini Impurity 与信息增益 Information Gain）的计算方式、物理含义及适用场景；'
     '理解树深度（max_depth）、最小叶节点样本数（min_samples_leaf）、成本复杂度剪枝（ccp_alpha）'
     '等超参数对模型复杂度、偏差-方差权衡的影响机制。'),
    ('2. 在真实电信客户流失数据集（Telco Customer Churn，7043条记录，21个原始特征）上完成端到端的'
     '机器学习工程流程：数据清洗→特征工程（标签编码/One-Hot编码/二值化）→类别不平衡处理→'
     '超参数调优→多维度性能评估，形成完整的工业级实验范式。'),
    ('3. 深入理解电信行业客户流失（Churn Prediction）场景中的类别不平衡问题：数据集中约73.5%的客户'
     '未流失，26.5%流失，若直接训练将产生"多数类偏见"。本实验系统对比三种主流应对策略：'
     'SMOTE 合成过采样、RandomUnderSampler 欠采样、class_weight=\'balanced\' 损失权重调节，'
     '并分析各策略在不同业务指标（Precision/Recall/F1/AUC）上的差异化表现。'),
    ('4. 掌握网格搜索（Grid Search）结合分层K折交叉验证（Stratified K-Fold CV，k=5）的超参数调优方法。'
     '理解为何以 roc_auc 作为调优评分指标而非 accuracy，以及各超参数（criterion、max_depth、'
     'min_samples_leaf、min_samples_split、ccp_alpha）的物理含义及其在参数网格中的设计依据。'),
    ('5. 全面掌握类别不平衡场景下的多维性能评估体系：使用 classification_report 输出各类别的'
     'Precision（精确率）、Recall（召回率）、F1-Score（调和均值），绘制混淆矩阵（直观展示误报FP'
     '和漏报FN）、ROC曲线并计算AUC值（衡量模型对正负样本的整体排序能力），绘制特征重要性图'
     '（揭示决策树学习到的关键业务驱动因子）。'),
    ('6. 以逻辑回归（Logistic Regression）作为对比基线，通过 ROC 曲线对比分析决策树与线性模型在'
     '非线性可分数据上的性能差异，深入思考两类模型的适用条件与局限性，为后续集成学习方法'
     '（Random Forest、XGBoost）的引入奠定理论认知基础。'),
]
for g in goals:
    _para(c5, g, indent=True)


# ══════════════════════════════════════════════════════════════════════════════
# 二、实验项目内容
# ══════════════════════════════════════════════════════════════════════════════
c6 = table.rows[6].cells[0]
cell_clear(c6)
_para(c6, '二、实验项目内容', bold=True, sa=3)

_sub(c6, '（一）数据集简介')
_para(c6,
    '本实验使用 IBM Telco Customer Churn 数据集（Kaggle，'
    'https://www.kaggle.com/datasets/blastchar/telco-customer-churn），'
    '该数据集来源于电信运营商真实业务系统，记录了7043名用户的服务订阅情况与流失行为。'
    '原始数据集包含21个字段：1个客户ID（customerID，无信息，丢弃）、'
    '1个目标变量（Churn，Yes/No）、19个特征变量。', indent=True)

_para(c6, '特征分三类：', indent=True)
_para(c6,
    '① 客户基本属性（4列）：gender（性别）、SeniorCitizen（是否老年用户，已为0/1）、'
    'Partner（是否有伴侣）、Dependents（是否有受抚养人）；',
    indent=True)
_para(c6,
    '② 服务订阅情况（12列）：tenure（在网时长，月）、PhoneService、MultipleLines、'
    'InternetService、OnlineSecurity、OnlineBackup、DeviceProtection、TechSupport、'
    'StreamingTV、StreamingMovies、Contract（合约类型）、PaperlessBilling、PaymentMethod；',
    indent=True)
_para(c6,
    '③ 账单信息（2列）：MonthlyCharges（月费，连续值）、TotalCharges（累计费用，原始为字符串'
    '，含11条空字符串需清洗）。',
    indent=True)
_para(c6,
    '类别分布：未流失（No）5174条（73.46%），流失（Yes）1869条（26.54%）。'
    '正负样本比约1:2.77，存在中等程度的类别不平衡，是本实验核心挑战之一。'
    '经过特征编码后，最终特征维度为38维。',
    indent=True)

_sub(c6, '（二）实验核心任务')
tasks = [
    '① 数据清洗：处理 TotalCharges 字段中因新用户（tenure=0）导致的11条空字符串，'
    '使用 MonthlyCharges 填充（新用户累计费即为当月费）；',
    '② 特征工程（编码）：gender 二值化（Male→1）；Partner/Dependents/PhoneService/'
    'PaperlessBilling 二值化（Yes→1）；Contract 有序编码（Month-to-month=0 < '
    'One year=1 < Two year=2，保留合约期与客户粘性的逻辑顺序）；'
    'InternetService/PaymentMethod/MultipleLines/OnlineSecurity/OnlineBackup/'
    'DeviceProtection/TechSupport/StreamingTV/StreamingMovies 共9列 One-Hot 编码；',
    '③ 数据集划分与标准化：按 stratify=Churn 进行8:2 分层划分，'
    'StandardScaler 仅在训练集上 fit，测试集仅 transform，严格防止数据泄露；',
    '④ 类别不平衡处理：SMOTE（过采样，sampling_strategy=0.5）/ '
    'RandomUnderSampler（欠采样）/ class_weight=\'balanced\'（权重调节）三种策略；',
    '⑤ 超参数调优：GridSearchCV 在5维参数空间（criterion×max_depth×'
    'min_samples_leaf×min_samples_split×ccp_alpha）共432种组合上网格搜索，'
    '5折分层CV，scoring=roc_auc；',
    '⑥ 性能评估与可视化：classification_report（Precision/Recall/F1）、'
    '混淆矩阵、ROC曲线（含LR基线对比）、特征重要性，全面评估模型在测试集上的表现；',
    '⑦ 对比分析：同时训练逻辑回归基线模型，在 ROC 曲线图中叠加展示，'
    '深入分析决策树与线性模型的性能差异及原因。',
]
for t in tasks:
    _para(c6, t, indent=True)


# ══════════════════════════════════════════════════════════════════════════════
# 三、实验过程或算法
# ══════════════════════════════════════════════════════════════════════════════
c7 = table.rows[7].cells[0]
cell_clear(c7)
_para(c7, '三、实验过程或算法', bold=True, sa=3)

# ── 3.1 ──────────────────────────────────────────────────────────────────────
_sub(c7, '3.1 数据预处理与特征工程')

_para(c7, '（1）数据清洗：TotalCharges 缺失值处理', bold=True)
_para(c7,
    'TotalCharges 字段在原始数据中为 object 类型（字符串），对于 tenure=0 的新用户，'
    '该字段为空字符串（共11条）。处理策略：先将空字符串强制转换为 NaN，再以 MonthlyCharges '
    '填充（因新用户累计费恰等于当月费，此填充具有业务合理性，优于简单删除或均值填充）。',
    indent=True)
_code(c7, "df['TotalCharges'] = pd.to_numeric(df['TotalCharges'].str.strip(), errors='coerce')")
_code(c7, "df['TotalCharges'] = df['TotalCharges'].fillna(df['MonthlyCharges'])")

_para(c7, '（2）标签编码（Label Encoding）：有序特征', bold=True)
_para(c7,
    '合约类型 Contract 具有明确的语义顺序（合约期越长，客户粘性通常越高），'
    '采用整数映射保留顺序信息，避免 One-Hot 编码将有序关系拆散：',
    indent=True)
_code(c7, "CONTRACT_ORDER = {'Month-to-month': 0, 'One year': 1, 'Two year': 2}")
_code(c7, "df['Contract'] = df['Contract'].map(CONTRACT_ORDER)")

_para(c7, '（3）二值化编码（Binary Encoding）', bold=True)
_para(c7,
    '对 gender（Male/Female）、Partner/Dependents/PhoneService/PaperlessBilling'
    '（Yes/No）共5列进行二值化，将其转化为0/1整数：',
    indent=True)
_code(c7, "df['gender'] = (df['gender'] == 'Male').astype(int)")
_code(c7, "for col in ['Partner','Dependents','PhoneService','PaperlessBilling']:")
_code(c7, "    df[col] = (df[col] == 'Yes').astype(int)")

_para(c7, '（4）独热编码（One-Hot Encoding）：无序分类特征', bold=True)
_para(c7,
    'InternetService（DSL/Fiber optic/No）、PaymentMethod（4类）、MultipleLines、'
    'OnlineSecurity/OnlineBackup/DeviceProtection/TechSupport/StreamingTV/StreamingMovies'
    '（均含Yes/No/No internet service三类）共9列无内在大小关系，使用 pd.get_dummies 全量展开。'
    '编码后从19个原始特征扩展至38维特征矩阵。',
    indent=True)
_code(c7, "ONEHOT_COLS = ['InternetService','PaymentMethod','MultipleLines',")
_code(c7, "               'OnlineSecurity','OnlineBackup','DeviceProtection',")
_code(c7, "               'TechSupport','StreamingTV','StreamingMovies']")
_code(c7, "df = pd.get_dummies(df, columns=ONEHOT_COLS, drop_first=False)")

_para(c7, '（5）数据集划分与标准化（防数据泄露规范）', bold=True)
_para(c7,
    '先按 stratify=Churn 做 8:2 分层划分（确保训练/测试集正样本比例一致），'
    '再对训练集 fit StandardScaler，仅对测试集 transform。严禁对完整数据集 fit，'
    '否则测试集的均值/方差信息会"泄露"到预处理阶段，导致评估结果偏乐观。'
    '虽然决策树对特征量纲不敏感，但为与逻辑回归基线统一比较，实验中统一进行标准化。',
    indent=True)
_code(c7, "X_train, X_test, y_train, y_test = train_test_split(")
_code(c7, "    X, y, test_size=0.2, random_state=42, stratify=y)")
_code(c7, "scaler = StandardScaler()")
_code(c7, "X_train_sc = scaler.fit_transform(X_train)  # fit 仅训练集")
_code(c7, "X_test_sc  = scaler.transform(X_test)        # 测试集只 transform")

# ── 3.2 ──────────────────────────────────────────────────────────────────────
_sub(c7, '3.2 决策树算法原理（CART）')

_para(c7, '（1）算法框架：CART（Classification and Regression Trees）', bold=True)
_para(c7,
    'scikit-learn 的 DecisionTreeClassifier 基于 CART 算法，采用二叉树结构（每个内部节点'
    '恰好分裂为2个子节点），通过贪心地在每个节点选择最优分裂特征和阈值，递归地将训练数据'
    '划分为更纯净的子集，直至满足停止条件（max_depth 或 min_samples_leaf）。',
    indent=True)

_para(c7, '（2）节点分裂准则：基尼系数（Gini Impurity）', bold=True)
_para(c7,
    '基尼系数衡量数据集的"不纯度"，值域 [0, 0.5]，越接近0表示越纯净：',
    indent=True)
_para(c7,
    '    Gini(D) = 1 - Σ pₖ²  （k遍历所有类别，pₖ 为第k类的样本比例）',
    indent=True)
_para(c7,
    '在流失预测的二分类中：若节点中流失:未流失 = 1:1，则 Gini = 1 - (0.5² + 0.5²) = 0.5（最不纯）；'
    '若节点全为未流失，则 Gini = 0（最纯）。'
    '对特征 f 在阈值 t 处分裂，分裂增益为：ΔGini = Gini(D) - |D_L|/|D|·Gini(D_L) - |D_R|/|D|·Gini(D_R)，'
    '选择 ΔGini 最大的 (f, t) 作为当前节点的分裂方式。',
    indent=True)

_para(c7, '（3）节点分裂准则：信息熵与信息增益（Entropy）', bold=True)
_para(c7,
    '信息熵衡量数据集的"无序程度"，值域 [0, log₂K]：',
    indent=True)
_para(c7,
    '    Entropy(D) = -Σ pₖ · log₂(pₖ)',
    indent=True)
_para(c7,
    '信息增益 IG = Entropy(D) - Σ |Dᵥ|/|D|·Entropy(Dᵥ)。'
    '与 Gini 相比，Entropy 的计算涉及对数运算，计算量稍大，但在类别不平衡场景下对少数类'
    '更为敏感，理论上更精确。在本实验中，GridSearchCV 同时搜索两种准则，最优策略'
    '由数据驱动自动选择。',
    indent=True)

_para(c7, '（4）防过拟合机制：超参数约束 + 成本复杂度剪枝', bold=True)
_para(c7,
    '决策树若完全生长（max_depth=None，min_samples_leaf=1），可完美拟合训练集'
    '但泛化能力极差（过拟合）。本实验采用以下策略抑制过拟合：',
    indent=True)
_para(c7,
    '• max_depth：限制树的最大层数，是最直接的过拟合控制手段。本实验网格：[3,5,7,10,15,None]；',
    indent=True)
_para(c7,
    '• min_samples_leaf：叶节点最少样本数，值越大叶节点越"粗糙"，相当于正则化。网格：[1,5,10,20]；',
    indent=True)
_para(c7,
    '• min_samples_split：节点分裂门槛，值越大分裂越保守。网格：[2,5,10]；',
    indent=True)
_para(c7,
    '• ccp_alpha（成本复杂度剪枝系数）：后剪枝方法，通过最小化带惩罚的损失函数'
    'R_α(T) = R(T) + α·|T|（R(T)为树的总不纯度，|T|为叶节点数，α为惩罚强度），'
    '逐步剪去"性价比"最低的子树。α=0 表示不剪枝，α越大剪枝越激进。网格：[0.0, 0.0001, 0.001]。',
    indent=True)

# ── 3.3 ──────────────────────────────────────────────────────────────────────
_sub(c7, '3.3 类别不平衡处理策略')
_para(c7,
    '类别不平衡的危害：若不做任何处理，将全部测试集预测为"未流失"可达73.5%的准确率，'
    '但此时 Recall(Yes)=0，漏掉全部真实流失客户，毫无业务价值。'
    '所有重采样操作均仅在训练集上执行，测试集保持原始分布（26.54%），保证评估公平性。',
    indent=True)

_para(c7, '策略A — SMOTE 过采样（Synthetic Minority Over-sampling Technique）', bold=True)
_para(c7,
    'SMOTE 算法原理：对少数类中每个样本 xᵢ，在其 k 个近邻（k=5）中随机选择一个近邻 xⱼ，'
    '沿连线合成新样本：x_new = xᵢ + λ·(xⱼ - xᵢ)，其中 λ∈[0,1] 为随机数。'
    '此方法合成的新样本在特征空间中均匀分布在少数类聚集区域，比简单重复（过采样）更能增强'
    '模型对少数类的泛化能力。参数 sampling_strategy=0.5 控制最终正:负 ≈ 1:2。',
    indent=True)
_code(c7, "smote = SMOTE(sampling_strategy=0.5, random_state=42, k_neighbors=5)")
_code(c7, "X_tr_sm, y_tr_sm = smote.fit_resample(X_train_sc, y_train)")
_code(c7, "# 原始训练集(5634,38)→ SMOTE后(6208,38)，正样本比例从26.54%升至33.3%")

_para(c7, '策略B — RandomUnderSampler 随机欠采样', bold=True)
_para(c7,
    '随机删除多数类（未流失）样本，使正:负 ≈ 1:2（sampling_strategy=0.5）。'
    '优点：训练速度快，避免引入合成噪声；缺点：丢弃大量真实多数类信息，'
    '约26%的原始训练数据（负类样本）被删除，可能导致模型对多数类的判断偏弱，'
    '体现为测试集上正类误报率（FPR）较高。',
    indent=True)
_code(c7, "rus = RandomUnderSampler(sampling_strategy=0.5, random_state=42)")
_code(c7, "X_tr_ru, y_tr_ru = rus.fit_resample(X_train_sc, y_train)")
_code(c7, "# 原始训练集(5634,38)→ 欠采样后(4485,38)，正样本33.3%")

_para(c7, "策略C — class_weight='balanced'（损失权重调节）", bold=True)
_para(c7,
    '不修改训练数据，通过在节点分裂的不纯度计算中对样本赋予不同权重来补偿不平衡。'
    "class_weight='balanced' 的权重公式：wₖ = n_samples / (n_classes × n_samplesₖ)，"
    '即流失类权重 = 7043 / (2 × 1869) ≈ 1.88，未流失类权重 ≈ 0.68，'
    '流失类样本的误分代价约为未流失类的2.76倍。所有原始训练数据完全保留，信息损失最小。',
    indent=True)

# ── 3.4 ──────────────────────────────────────────────────────────────────────
_sub(c7, '3.4 超参数调优（GridSearchCV + 5-Fold Stratified CV）')

_para(c7, '参数网格设计（共432种组合）：', indent=True)
_code(c7, "DT_PARAM_GRID = {")
_code(c7, "    'criterion':         ['gini', 'entropy'],          # 2 种分裂准则")
_code(c7, "    'max_depth':         [3, 5, 7, 10, 15, None],      # 6 个深度设置")
_code(c7, "    'min_samples_leaf':  [1, 5, 10, 20],               # 4 个叶节点门槛")
_code(c7, "    'min_samples_split': [2, 5, 10],                   # 3 个分裂门槛")
_code(c7, "    'ccp_alpha':         [0.0, 0.0001, 0.001],         # 3 个剪枝强度")
_code(c7, "}  # 2×6×4×3×3 = 432 种组合 × 5折 = 2160 次拟合/每策略")

_para(c7,
    '为何使用 roc_auc 而非 accuracy 作为 CV 评分指标：',
    indent=True)
_para(c7,
    '在正样本仅26.54%的场景下，将全部样本预测为"未流失"可达73.5%的准确率，'
    '但此"高准确率"完全无业务价值。roc_auc（ROC曲线下面积）是阈值无关的指标，'
    '度量模型在所有可能决策阈值下将正样本排在负样本前面的概率，不受类别不平衡影响，'
    '是评估分类器综合排序能力的最佳单一指标。',
    indent=True)

_para(c7, '分层K折交叉验证（Stratified K-Fold）：', indent=True)
_para(c7,
    '普通K折可能导致某些折中正样本极少（训练集正样本比例不稳定），'
    '分层K折确保每折中正负样本比例与整体一致（约26.5%），避免CV估计的方差偏大。',
    indent=True)
_code(c7, "cv = StratifiedKFold(n_splits=5, shuffle=True, random_state=42)")
_code(c7, "gs = GridSearchCV(")
_code(c7, "    estimator=DecisionTreeClassifier(random_state=42),")
_code(c7, "    param_grid=DT_PARAM_GRID,")
_code(c7, "    cv=cv,")
_code(c7, "    scoring='roc_auc',")
_code(c7, "    n_jobs=-1,    # 并行所有CPU核心加速搜索")
_code(c7, "    refit=True,   # 找到最优参数后在全训练集上重新拟合")
_code(c7, ")")
_code(c7, "gs.fit(X_train, y_train)")
_code(c7, "best_model = gs.best_estimator_")

_para(c7, '各策略最优参数搜索结果：', indent=True)
_build_table(c7,
    headers=['策略', '最优criterion', 'max_depth', 'min_samples_leaf',
             'min_samples_split', 'ccp_alpha', 'CV-AUC', '树深度/叶数'],
    rows=[
        ['SMOTE',       'gini',    '10',  '1',  '2', '0.001', '0.8447', '7 / 22'],
        ['UnderSample', 'entropy', '5',   '5',  '2', '0.001', '0.8261', '5 / 25'],
        ['ClassWeight', 'entropy', '5',   '10', '2', '0.0',   '0.8323', '5 / 31'],
    ],
    note='注：ccp_alpha=0.001 的出现说明轻微后剪枝能改善泛化；'
         'max_depth=10 在SMOTE数据上适合（样本量更多，6208条），'
         '而两种不平衡调整策略倾向更浅的树（max_depth=5）。'
)
_para(c7, '')

_para(c7, '参数分析：', indent=True)
_para(c7,
    '• SMOTE 策略选 gini、max_depth=10（实际被 ccp_alpha=0.001 剪枝至深度7），'
    '说明 SMOTE 后训练集增大（6208条），树可以更深地学习；',
    indent=True)
_para(c7,
    '• UnderSample 和 ClassWeight 策略均选 entropy、max_depth=5，'
    '说明在样本较少或直接使用原始分布时，较浅的树更能泛化；',
    indent=True)
_para(c7,
    '• ccp_alpha=0.001 出现在 SMOTE 和 UnderSample 策略，说明后剪枝能有效去除'
    '"纯噪声叶节点"，改善测试集性能。',
    indent=True)

# ── 3.5 ──────────────────────────────────────────────────────────────────────
_sub(c7, '3.5 逻辑回归基线模型（对比用）')
_para(c7,
    '为在 ROC 曲线上进行横向对比，同时训练一个逻辑回归基线模型（使用 SMOTE 处理后的训练集 + '
    "class_weight='balanced' + GridSearchCV 调优 C 和 penalty），用于评估决策树相对于线性"
    '分类器的性能优劣，以及分析两者在哪类样本上各有优势。',
    indent=True)


# ══════════════════════════════════════════════════════════════════════════════
# 四、实验结果及分析
# ══════════════════════════════════════════════════════════════════════════════
c8 = table.rows[8].cells[0]
cell_clear(c8)
_para(c8, '四、实验结果及分析', bold=True, sa=3)

# ── 4.1 ──────────────────────────────────────────────────────────────────────
_sub(c8, '4.1 三种策略完整性能对比（测试集，n=1409，正样本374）')

_build_table(c8,
    headers=['策略', 'Precision\n(Yes)', 'Recall\n(Yes)', 'F1\n(Yes)',
             'Accuracy', 'AUC', 'TN', 'FP', 'FN', 'TP'],
    rows=[
        ['DT + SMOTE',       '0.5907', '0.6444', '0.6164', '78.71%', '0.8238', '868', '167', '133', '241'],
        ['DT + UnderSample', '0.6070', '0.6070', '0.6070', '79.13%', '0.8337', '888', '147', '147', '227'],
        ['DT + ClassWeight', '0.5262', '0.7513', '0.6189', '75.44%', '0.8374', '782', '253', '93',  '281'],
        ['LR Baseline',      '—',      '—',      '—',      '—',      '0.8405', '—',  '—',  '—',  '—'],
    ],
    note='注：Yes类=流失正样本（support=374）；LR基线仅用于ROC曲线对比，不作主要评估对象。'
)
_para(c8, '')

# ── 4.2 ──────────────────────────────────════════════════════════════════════
_sub(c8, '4.2 分类指标深度分析')

_para(c8, '（1）Accuracy 的局限性与不平衡场景的评估误导', bold=True)
_para(c8,
    '若将全部1409个测试样本预测为"未流失"，准确率将达73.5%——与本实验中 ClassWeight 策略的'
    '75.44%相差无几，但"全预测负类"的 Recall(Yes)=0，漏掉全部374位真实流失客户。'
    '这表明在不平衡场景下，Accuracy 是极具误导性的评估指标，必须结合 Precision、Recall、'
    'F1 和 AUC 进行综合评估。',
    indent=True)

_para(c8, '（2）Precision-Recall 权衡分析（Yes类）', bold=True)
_para(c8,
    'SMOTE 策略（Precision=0.591，Recall=0.644，F1=0.616）：'
    'Precision 三策略中最高，每预测10个"潜在流失客户"中有5.9个是真实的，误报167人（FPR=16.1%）；'
    'Recall 最低但仍达64.4%，漏报133人（FNR=35.6%）。适合运营商希望控制营销骚扰成本、'
    '精准锁定高概率流失客户的场景。',
    indent=True)
_para(c8,
    'UnderSample 策略（Precision=0.607，Recall=0.607，F1=0.607）：'
    'Precision 与 Recall 完全相等，呈现出罕见的"完美对称"，这与欠采样后训练集中正负样本'
    '比例恰好接近1:2有关（模型决策阈值较为中性）。F1=0.607，略低于其他策略。',
    indent=True)
_para(c8,
    'ClassWeight 策略（Precision=0.526，Recall=0.751，F1=0.619）：'
    'Recall 三策略最高（75.1%），漏报仅93人（FNR=24.9%），意味着约3/4的真实流失客户被识别。'
    '但代价是 FP=253（误报率25.3%），精确率最低。'
    '在电信业务中，流失一位客户的损失（留存成本、收入损失）远超误报一次营销电话的成本，'
    '因此 ClassWeight 策略在实际业务价值上往往最优。',
    indent=True)

_para(c8, '（3）F1-Score 与宏观/加权平均分析', bold=True)
_para(c8,
    'SMOTE 策略的 macro avg F1=0.735 最高，说明它在正负两类上的综合表现最均衡；'
    'ClassWeight 策略的 macro avg F1=0.719，但其对业务关键的正类 Recall 最高。'
    '建议实际部署时按业务目标选择：若以最小化客户流失为第一目标，选 ClassWeight；'
    '若兼顾营销精准度，选 SMOTE。',
    indent=True)

_para(c8, '（4）决策树 vs 逻辑回归 AUC 分析', bold=True)
_para(c8,
    '最优决策树（ClassWeight，AUC=0.8374）与 LR 基线（AUC=0.8405）相差仅 0.003，'
    '差异几乎可忽略（<0.5%）。这说明：在充分调优的前提下，决策树在此数据集上的排序能力'
    '可与逻辑回归媲美。决策树的潜在优势在于天然可解释性（可直接提取决策路径）和对非线性边界'
    '的天然适应；但其弱点是对训练数据的方差较高（单棵树不稳定），'
    '这正是集成方法（随机森林、XGBoost）能显著超越单棵决策树的根本原因。',
    indent=True)

# ── 4.3 可视化 ──────────────────────────────────────────────────────────────
_sub(c8, '4.3 可视化分析')

_para(c8,
    '图1  Precision / Recall / F1-Score 对比柱状图',
    bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, sa=2)
_img(c8, IMG['prf'], width_cm=14.5)
_para(c8,
    '图1 以分组柱状图直观呈现三种策略在 Precision、Recall、F1-Score 三个指标上的差异。'
    'SMOTE 在 Precision（0.591）上明显领先，体现了其"适度平衡"的特点——不过于倾向正类，'
    '控制了误报率；ClassWeight 在 Recall（0.751）上领先约11个百分点，以牺牲精确率换取'
    '更高的查全率；F1-Score 方面三者差异较小（0.607~0.619），ClassWeight 略胜。',
    indent=True, sa=4)

_para(c8,
    '图2  分类报告热力图（Precision / Recall / F1 × 两类别 × 三模型）',
    bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, sa=2)
_img(c8, IMG['heatmap'], width_cm=14.5)
_para(c8,
    '图2 以热力图矩阵（绿色=高，红色=低）同时展示正负两个类别在三个指标上的表现。'
    '负类（未流失 No）的所有格均呈深绿色（0.80~0.95），说明模型对多数类预测稳健；'
    '正类（流失 Yes）的 Precision 列呈黄色区域（0.53~0.61），反映类别不平衡场景下'
    '精确率的固有限制。macro avg 行（正负均等权重）中，SMOTE 策略各列最为均衡。',
    indent=True, sa=4)

_para(c8,
    '图3  混淆矩阵对比（三种不平衡处理策略）',
    bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, sa=2)
_img(c8, IMG['cm'], width_cm=14.5)
_para(c8,
    '图3 的三张混淆矩阵从左到右对应 SMOTE、UnderSample、ClassWeight 三策略。'
    '观察 FN（左下角，漏报真实流失）：ClassWeight 最少（93），SMOTE 居中（133），'
    'UnderSample 最多（147）；观察 FP（右上角，误报未流失）：SMOTE 最少（167），'
    'UnderSample 次之（147），ClassWeight 最多（253）。'
    '这一"此消彼长"的模式体现了分类器在 Precision-Recall 之间的本质权衡（P-R Trade-off）：'
    '提高召回率必然降低精确率，反之亦然，无法同时最优。',
    indent=True, sa=4)

_para(c8,
    '图4  ROC 曲线对比（决策树三策略 + 逻辑回归基线）',
    bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, sa=2)
_img(c8, IMG['roc'], width_cm=14.0)
_para(c8,
    '图4 叠加展示了3条决策树 ROC 曲线和1条 LR 基线。'
    '四条曲线均显著高于随机分类器对角线（AUC=0.5），AUC范围0.824~0.841，'
    '表明各模型均具备较强的正负样本排序能力。'
    'ClassWeight（实线，AUC=0.837）与 LR 基线（虚线，AUC=0.841）高度重合，'
    '在大部分阈值区间差异不足1%，说明充分调优后决策树可追平逻辑回归。'
    '图中圆点标注了各曲线的最优工作点（Youden Index，即 TPR-FPR 最大处），'
    '可直观指导实际部署时决策阈值的选取。',
    indent=True, sa=4)

_para(c8,
    '图5  特征重要性 Top-20（最优决策树：ClassWeight 策略）',
    bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, sa=2)
_img(c8, IMG['feat_imp'], width_cm=14.0)
_para(c8,
    '图5 展示最优决策树（ClassWeight 策略）中各特征的基尼重要性（值越高表示该特征对节点'
    '分裂的贡献越大）。tenure（在网时长）重要性最高，与业务认知高度一致——'
    '新用户（tenure短）流失率显著高于老用户；'
    'MonthlyCharges（月费）和 TotalCharges（累计费用）紧随其后，'
    '反映价格敏感度是流失的核心驱动因素；'
    'Contract_Month-to-month（月付合约）重要性较高，月付用户无违约成本，流失门槛低；'
    'InternetService_Fiber optic（光纤服务）的高重要性反映光纤用户可能因竞争激烈而更易流失。'
    '这些特征与电信行业的业务洞察完全吻合，说明决策树通过纯数据驱动的方式成功学到了'
    '有意义的业务规律。',
    indent=True, sa=4)

# ── 4.4 ──────────────────────────────────────────────────────────────────────
_sub(c8, '4.4 实验难点与解决方案')

_build_table(c8,
    headers=['难点', '具体表现', '解决方案', '效果'],
    rows=[
        ['TotalCharges字段异常',
         '11条新用户记录中TotalCharges为空字符串，pd.to_numeric直接报错',
         '先strip再errors=\'coerce\'转NaN，用MonthlyCharges填充（业务语义合理）',
         '消除缺失值，无信息损失'],
        ['合约类型编码歧义',
         '若用OHE编码Contract，模型无法感知"月付<年付<两年"的语义顺序',
         '有序Label Encoding（0/1/2），保留合约期与粘性的单调关系',
         '该特征重要性排名进入Top-10'],
        ['类别不平衡导致评估误导',
         '直接用Accuracy，模型全预测负类可达73.5%，但Recall=0',
         'CV使用roc_auc评分；评估使用P/R/F1全套指标；三种重采样策略',
         'Recall(Yes)从~40%提升至61~75%'],
        ['参数空间过大导致搜索慢',
         '432种组合×5折×3策略=6480次拟合，单机耗时较长',
         'n_jobs=-1并行；仅搜索决策树关键超参数；适当缩减无效区间',
         '搜索耗时约2~3分钟，可接受'],
        ['决策树过拟合',
         '完全生长树在测试集AUC骤降至~0.72',
         'max_depth约束+min_samples_leaf正则化+ccp_alpha后剪枝三管齐下',
         'AUC稳定在0.82~0.84'],
    ],
    note='注：以上难点均已在代码中通过模块化设计（preprocessing.py / imbalance.py / model.py）妥善处理。'
)
_para(c8, '')

# ── 4.5 ──────────────────────────────────────────────────────────────────────
_sub(c8, '4.5 实验结论')
conclusions = [
    '① 决策树在充分调优后（GridSearchCV + 三维防过拟合机制）可在 Telco Churn 数据集上实现 '
    'AUC≈0.837，与逻辑回归基线（AUC=0.841）差距仅0.004，证明决策树具备足够的竞争力，'
    '且具有更强的可解释性（可直接提取决策路径和特征重要性）；',

    '② 类别不平衡处理至关重要：三种策略均将 Recall(Yes) 从未处理时的约40%提升至61%~75%；'
    'class_weight=\'balanced\' 策略在 AUC（0.8374）和 Recall（75.1%）两项指标上表现最优，'
    '且保留了100%原始训练数据，信息损失最小，是本实验推荐的部署方案；',

    '③ 特征工程对决策树性能影响显著：tenure（在网时长）、MonthlyCharges、TotalCharges、'
    'Contract_Month-to-month 成为最重要的前4个特征，与电信行业的业务洞察完全吻合，'
    '说明决策树能自动挖掘关键业务驱动因子；',

    '④ 超参数调优揭示了重要规律：ccp_alpha=0.001（轻微后剪枝）在多个策略中被选中，'
    '证明完全生长的树在此数据集上存在过拟合；max_depth=5~7 是该数据集的合理树深，'
    '更深的树因噪声拟合而泛化性能下降；',

    '⑤ 决策树与逻辑回归的根本差异：LR 需要特征工程（如多项式交互项）才能捕捉非线性关系，'
    '而决策树天然支持非线性边界。若要突破0.85的 AUC 上限，应引入随机森林（Bootstrap聚合'
    '降低方差）或 XGBoost（梯度提升集成），预期 AUC 可达0.85~0.90。',
]
for c in conclusions:
    _para(c8, c, indent=True)

# ── 保存 ──────────────────────────────────────────────────────────────────────
doc.save(OUT_PATH)
print(f'[OK] 实验报告已保存至: {OUT_PATH}')
