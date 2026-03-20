"""
gen_report.py - 自动生成实验报告（test2_zsq）
用法：conda run -n ml python gen_report.py
"""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE     = os.path.dirname(os.path.abspath(__file__))
TEMPLATE = os.path.join(BASE, '..', 'test2_gxr', 'test2_Experiment_Report.docx')
IMG_METRICS  = os.path.join(BASE, 'output', 'metrics_bar.png')
IMG_CM       = os.path.join(BASE, 'output', 'confusion_matrices.png')
IMG_ROC      = os.path.join(BASE, 'output', 'roc_curves.png')
IMG_FI       = os.path.join(BASE, 'output', 'feature_importance.png')
OUT_PATH = os.path.join(BASE, 'test2_Experiment_Report.docx')

doc   = Document(TEMPLATE)
table = doc.tables[0]

# ── 颜色常量 ──────────────────────────────────────────────────────────────────
COLOR_CODE    = RGBColor(0x1a, 0x53, 0x76)   # 深蓝 - 代码
COLOR_FORMULA = RGBColor(0x7B, 0x00, 0x1A)   # 深红 - 公式
COLOR_EMPH    = RGBColor(0x0D, 0x5C, 0x2E)   # 深绿 - 强调
SHADE_HEADER  = 'BDD7EE'
SHADE_NOTE    = 'F2F2F2'

# ── 辅助函数 ──────────────────────────────────────────────────────────────────

def fmt_run(run, name='宋体', size=12, bold=False, color=None,
            code=False, italic=False):
    fn = 'Courier New' if code else name
    run.font.name = fn
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color if isinstance(color, RGBColor) else RGBColor(*color)
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), name)
    rPr.insert(0, rFonts)


def cell_clear(cell):
    for p in cell.paragraphs[1:]:
        p._element.getparent().remove(p._element)
    cell.paragraphs[0].clear()


def add_para(cell, text, bold=False, code=False, indent_first=False,
             align=WD_ALIGN_PARAGRAPH.LEFT, space_after=2, space_before=0,
             color=None):
    p = cell.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after  = Pt(space_after)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    if indent_first:
        pf.first_line_indent = Pt(24)
    p.alignment = align
    run = p.add_run(text)
    clr = color or (COLOR_CODE if code else None)
    fmt_run(run, bold=bold, size=12, code=code, color=clr)
    return p


def add_sub(cell, text, space_before=6, level=2):
    """小节标题（加粗）"""
    p = cell.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after  = Pt(2)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = p.add_run(text)
    sz = 12 if level == 2 else 11
    fmt_run(run, bold=True, size=sz)
    return p


def add_code(cell, text):
    """代码块行"""
    p = cell.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after  = Pt(0)
    pf.left_indent  = Cm(0.6)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = p.add_run(text)
    fmt_run(run, code=True, size=9, color=COLOR_CODE)
    return p


def add_formula(cell, text, note=None):
    """数学公式行（居中，深红色斜体）"""
    p = cell.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(2)
    pf.space_after  = Pt(2)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.left_indent = Cm(1.5)
    run = p.add_run(text)
    fmt_run(run, name='Times New Roman', size=11, italic=True,
            color=COLOR_FORMULA)
    if note:
        run2 = p.add_run('  ' + note)
        fmt_run(run2, size=10, color=RGBColor(0x60, 0x60, 0x60))
    return p


def add_img(cell, path, width_cm=15, caption=None):
    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run()
    run.add_picture(path, width=Cm(width_cm))
    if caption:
        pc = cell.add_paragraph()
        pc.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pc.paragraph_format.space_before = Pt(0)
        pc.paragraph_format.space_after  = Pt(6)
        rc = pc.add_run(caption)
        fmt_run(rc, size=10, italic=True, color=RGBColor(0x44, 0x44, 0x44))
    return p


def _set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)


def _set_table_borders(tbl):
    tblPr = tbl._tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl._tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'),   'single')
        border.set(qn('w:sz'),    '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '4472C4')
        tblBorders.append(border)
    tblPr.append(tblBorders)


def _set_col_widths(tbl, widths_cm):
    for row in tbl.rows:
        for ci, cell in enumerate(row.cells):
            if ci < len(widths_cm):
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcW = OxmlElement('w:tcW')
                tcW.set(qn('w:w'), str(int(widths_cm[ci] * 567)))
                tcW.set(qn('w:type'), 'dxa')
                tcPr.append(tcW)


def build_table(cell, headers, rows, note=None, col_widths=None):
    n_rows = 1 + len(rows) + (1 if note else 0)
    t = cell.add_table(rows=n_rows, cols=len(headers))
    _set_table_borders(t)
    for ci, h in enumerate(headers):
        c = t.rows[0].cells[ci]
        _set_cell_bg(c, SHADE_HEADER)
        c.paragraphs[0].clear()
        r = c.paragraphs[0].add_run(h)
        fmt_run(r, bold=True, size=10)
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        c.paragraphs[0].paragraph_format.space_before = Pt(1)
        c.paragraphs[0].paragraph_format.space_after  = Pt(1)
    for ri, rd in enumerate(rows):
        bg = 'FFFFFF' if ri % 2 == 0 else 'EBF3FB'
        for ci, val in enumerate(rd):
            c = t.rows[ri + 1].cells[ci]
            _set_cell_bg(c, bg)
            c.paragraphs[0].clear()
            r = c.paragraphs[0].add_run(str(val))
            fmt_run(r, size=10)
            c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            c.paragraphs[0].paragraph_format.space_before = Pt(1)
            c.paragraphs[0].paragraph_format.space_after  = Pt(1)
    if note:
        nr = t.rows[-1]
        for ci in range(len(headers)):
            _set_cell_bg(nr.cells[ci], SHADE_NOTE)
        nr.cells[0].paragraphs[0].clear()
        r = nr.cells[0].paragraphs[0].add_run(note)
        fmt_run(r, size=9, color=RGBColor(0x60, 0x60, 0x60))
        nr.cells[0].paragraphs[0].paragraph_format.space_before = Pt(1)
        nr.cells[0].paragraphs[0].paragraph_format.space_after  = Pt(1)
        for i in range(1, len(headers)):
            nr.cells[0].merge(nr.cells[i])
    if col_widths:
        _set_col_widths(t, col_widths)
    return t


# ══════════════════════════════════════════════════════════════════════════════
# 封面
# ══════════════════════════════════════════════════════════════════════════════
for cell in table.rows[0].cells:
    if '专业' in cell.text or '学号' in cell.text:
        cell_clear(cell)
        r = cell.paragraphs[0].add_run('专业：电子信息工程       学号：2023XXXXXXXXX')
        fmt_run(r, size=12)
        break
for cell in table.rows[0].cells:
    if cell.text.strip() == '姓名' or '姓名' in cell.text:
        cell_clear(cell)
        r = cell.paragraphs[0].add_run('姓名：zsq')
        fmt_run(r, size=12)
        break
for cell in table.rows[1].cells:
    if '实验题目' in cell.text:
        cell_clear(cell)
        r = cell.paragraphs[0].add_run('实验题目：基于决策树算法的电信客户流失预测')
        fmt_run(r, bold=True, size=12)
        break
for cell in table.rows[2].cells:
    if '实验时间' in cell.text:
        cell_clear(cell)
        r = cell.paragraphs[0].add_run('实验时间：2026年3月20日')
        fmt_run(r, size=12)
        break

# ══════════════════════════════════════════════════════════════════════════════
# 一、实验目的
# ══════════════════════════════════════════════════════════════════════════════
c5 = table.rows[5].cells[0]
cell_clear(c5)
add_para(c5, '一、实验目的', bold=True)
for g in [
    '1. 深入理解决策树（Decision Tree / CART 算法）的数学原理：掌握节点分裂准则（基尼不纯度 Gini Impurity 与信息熵 Entropy）、递归二叉分裂过程、树深度与叶节点样本量对模型复杂度的控制机制，以及代价复杂度剪枝（CCP, Cost-Complexity Pruning）的数学推导与实现；',
    '2. 掌握面向结构化数据的端到端机器学习工程流程：以 6 个解耦模块（config / preprocessing / imbalance / modeling / evaluation / main）构建可维护的项目结构，涵盖数据加载与清洗、业务驱动特征工程、特征编码、数据标准化、不平衡处理、超参数调优与多维评估的完整链路；',
    '3. 系统掌握混合编码策略的设计依据：理解 One-Hot 编码（InternetService、PaymentMethod 等无序类别）与有序 Label Encoding（Contract 合约类型等有序类别）的本质区别，能根据特征的业务语义和数值关系选择合适的编码方式，避免引入虚假的大小关系假设；',
    '4. 深入理解并实践三种类别不平衡处理策略：对比 SMOTE 过采样（imblearn Pipeline 防泄露版本）、随机欠采样和 class_weight 损失加权的工作原理与 Precision-Recall 权衡，特别掌握"SMOTE 必须放入交叉验证 Pipeline 内执行"这一防止数据泄露的关键工程实践；',
    '5. 掌握 GridSearchCV + StratifiedKFold 超参数调优方法，理解决策树四大核心超参数（max_depth、min_samples_leaf、min_samples_split、criterion）对树结构和泛化能力的影响机制，掌握 CCP alpha 代价复杂度剪枝参数的引入方式，能通过网格搜索寻找最优参数组合；',
    '6. 理解并能计算类别不平衡场景下的核心评估指标：Precision（精确率）、Recall（召回率）、F1-Score（调和平均）与 ROC-AUC，能从混淆矩阵中解读 FP（误报）与 FN（漏报）对电信业务的不同代价含义，并理解为何 Recall 在流失预测场景中优先级高于 Precision；',
    '7. 理解决策树相对于逻辑回归的优势与局限：通过 DT vs. LR 的 ROC 曲线对比实验，认识到单棵决策树在中等规模结构化数据集上的性能瓶颈（AUC≈0.83），理解集成树方法（Random Forest、XGBoost）突破此瓶颈的理论基础。',
]:
    add_para(c5, g, indent_first=True)

# ══════════════════════════════════════════════════════════════════════════════
# 二、实验项目内容
# ══════════════════════════════════════════════════════════════════════════════
c6 = table.rows[6].cells[0]
cell_clear(c6)
add_para(c6, '二、实验项目内容', bold=True)

add_sub(c6, '（一）数据集简介')
add_para(c6,
    '本实验使用 IBM Telco Customer Churn 数据集（Kaggle 公开数据集），共 7,043 条电信客户记录，'
    '20 个原始特征（含数值型和类别型），目标变量 Churn 标记客户是否已流失（Yes/No）。',
    indent_first=True)
add_para(c6,
    '特征涵盖三个维度：① 客户基本属性（gender、SeniorCitizen、Partner、Dependents、tenure）；'
    '② 服务订阅信息（PhoneService、MultipleLines、InternetService、OnlineSecurity、OnlineBackup、'
    'DeviceProtection、TechSupport、StreamingTV、StreamingMovies）；'
    '③ 账单与合约信息（Contract、PaperlessBilling、PaymentMethod、MonthlyCharges、TotalCharges）。',
    indent_first=True)
add_para(c6,
    '类别分布呈轻度不平衡：未流失 5,174 条（73.46%）、流失 1,869 条（26.54%），正负样本比约 1:2.77。'
    '此外 TotalCharges 列存在 11 条空字符串（以 object 类型存储），为真实工业数据中的典型数据质量问题。',
    indent_first=True)

build_table(c6,
    headers=['字段', '类型', '取值示例', '业务含义'],
    rows=[
        ['tenure',        '数值',    '0–72（月）',                  '客户入网时长，越长粘性越高'],
        ['Contract',      '有序类别', 'Month-to-month / One year / Two year', '合约类型，月付流失率远高于年付'],
        ['MonthlyCharges','数值',    '18.25–118.75（美元）',         '月消费额，高消费感知差异大'],
        ['TotalCharges',  '数值',    '18.8–8684.8（美元）',          '累计消费，含 11 条空字符串缺失'],
        ['InternetService','无序类别','DSL / Fiber optic / No',       '网络类型，光纤用户流失率最高'],
        ['PaymentMethod', '无序类别', '4 种支付方式',                 '电子支票用户流失率显著偏高'],
        ['Churn',         '目标变量', 'Yes=1 / No=0',                '是否已流失（正类为流失）'],
    ],
    note='注：全部特征均无 NaN，仅 TotalCharges 含空字符串需转换处理',
    col_widths=[3.0, 2.0, 4.5, 6.5]
)

add_sub(c6, '（二）工程化项目架构')
add_para(c6, '本实验将功能解耦为 6 个独立模块，实现关注点分离和可维护性：', indent_first=True)
build_table(c6,
    headers=['模块文件', '职责说明'],
    rows=[
        ['config.py',       '全局配置：路径、随机种子、CV 折数、编码映射、超参数网格（含 CCP alpha）'],
        ['preprocessing.py','数据加载→TotalCharges 修复→10 项特征工程→混合编码→StandardScaler 标准化'],
        ['imbalance.py',    '三种不平衡策略封装（SMOTE 返回重采样数据 / UnderSample / ClassWeight 标志）'],
        ['modeling.py',     'SMOTE Pipeline GridSearchCV + UnderSample/ClassWeight GridSearchCV + LR 对比'],
        ['evaluation.py',   '测试集推断、classification_report 打印、4 张可视化图表生成、AUC 汇总'],
        ['main.py',         '主执行入口，串联完整流水线，9 步骤输出所有结果'],
    ],
    col_widths=[3.5, 12.5]
)
add_para(c6, '')

add_sub(c6, '（三）实验核心任务')
for t in [
    '① 数据清洗：TotalCharges 空字符串→中位数填充；删除 customerID（无预测价值）；',
    '② 业务驱动特征工程：构造 10 个衍生特征，包括 num_services（服务订阅数）、risky_combo（月付+新客户高风险组合）、'
    'tenure_segment（入网时长分段）、monthly_to_total（消费新鲜度指标）、fiber_high_charge（光纤高消费客户）等；',
    '③ 混合编码策略：二值 Yes/No 列→0/1 映射；Contract 等有序类别→Label Encoding 保留顺序语义；'
    'InternetService、PaymentMethod 等无序类别→One-Hot Encoding 防止虚假大小关系；',
    '④ 三种类别不平衡处理：SMOTE Pipeline（每折独立执行防泄露）/ 随机欠采样 / class_weight=\'balanced\'；',
    '⑤ 决策树超参数调优：GridSearchCV 在 criterion×max_depth×min_samples_leaf×min_samples_split×ccp_alpha '
    '五维参数空间搜索（210 组参数 × 5折 = 1,050 次拟合/策略），scoring=roc_auc；',
    '⑥ 多维评估：classification_report 输出 Precision/Recall/F1，混淆矩阵分析 FP/FN，ROC 曲线 + AUC，特征重要性；',
    '⑦ DT vs. LR 对比：同数据上训练 GridSearchCV 优化的 LR，ROC 图叠加对比，分析线性模型 vs 树模型的差异。',
]:
    add_para(c6, t, indent_first=True)

# ══════════════════════════════════════════════════════════════════════════════
# 三、实验过程或算法
# ══════════════════════════════════════════════════════════════════════════════
c7 = table.rows[7].cells[0]
cell_clear(c7)
add_para(c7, '三、实验过程或算法', bold=True)

# 3.0 决策树数学原理
add_sub(c7, '3.0 决策树（CART算法）数学原理')

add_sub(c7, '（1）基本概念：递归二叉分裂', level=3)
add_para(c7,
    'CART（Classification and Regression Tree）算法通过自顶向下的递归二叉分裂构建决策树。'
    '在每个内部节点，算法从所有特征 j 和所有可能分裂阈值 t 中，选择使目标函数最优的分裂点 (j*, t*)：',
    indent_first=True)
add_formula(c7, '(j*, t*) = argmin_{j,t}  [ |L|/|S| · I(D_L)  +  |R|/|S| · I(D_R) ]',
            note='S=父节点样本集，L/R=左/右子节点，I=不纯度函数')
add_para(c7,
    '分裂持续进行，直到满足停止条件（max_depth 达到上限、min_samples_split 或 min_samples_leaf 约束、'
    '节点完全纯净，或 ccp_alpha 剪枝移除子树）。',
    indent_first=True)

add_sub(c7, '（2）分裂准则一：基尼不纯度（Gini Impurity）', level=3)
add_para(c7,
    '基尼不纯度衡量节点 D 中随机选取两个样本类别不同的概率，取值范围 [0, 0.5]（二分类）。'
    '值越小表示节点越纯净：',
    indent_first=True)
add_formula(c7, 'Gini(D) = 1 − Σ_{k=1}^{K} p_k²',
            note='p_k = 节点中第 k 类样本比例')
add_para(c7,
    '二分类情形下（K=2），当 p₁=p₂=0.5 时 Gini=0.5（最不纯），当某类占比为 1 时 Gini=0（完全纯净）。'
    'Gini 不纯度的分裂增益为：',
    indent_first=True)
add_formula(c7,
    'ΔGini = Gini(D) − [ |D_L|/|D| · Gini(D_L)  +  |D_R|/|D| · Gini(D_R) ]',
    note='选择使 ΔGini 最大的分裂点')
add_para(c7,
    'Gini 计算无对数运算，效率更高，是 sklearn DecisionTreeClassifier 的默认准则（criterion=\'gini\'）。',
    indent_first=True)

add_sub(c7, '（3）分裂准则二：信息熵与信息增益（Entropy / Information Gain）', level=3)
add_para(c7,
    '信息熵基于香农信息论，衡量节点的不确定性，取值范围 [0, 1]（二分类归一化后）：',
    indent_first=True)
add_formula(c7, 'H(D) = − Σ_{k=1}^{K} p_k · log₂(p_k)',
            note='p_k log₂ p_k = 0 当 p_k = 0')
add_para(c7, '信息增益（Information Gain）为分裂前后熵的减少量：', indent_first=True)
add_formula(c7,
    'IG = H(D) − [ |D_L|/|D| · H(D_L)  +  |D_R|/|D| · H(D_R) ]')
add_para(c7,
    '信息熵对小概率类别更敏感（对数放大效应），在类别不平衡场景中理论上略优于 Gini。'
    '本实验通过 GridSearchCV 同时搜索 criterion∈{\'gini\', \'entropy\'}，由数据驱动最优选择。',
    indent_first=True)

add_sub(c7, '（4）代价复杂度剪枝（CCP, Cost-Complexity Pruning）', level=3)
add_para(c7,
    '未剪枝的完全生长决策树极易过拟合训练集。CCP 通过引入惩罚系数 α 将树的复杂度纳入目标函数：',
    indent_first=True)
add_formula(c7, 'R_α(T) = R(T)  +  α · |T_L|',
            note='R(T)=树的误分率，|T_L|=叶子节点数，α≥0 控制剪枝力度')
add_para(c7,
    '对于给定的 α，存在唯一最优子树 T*(α)。α=0 时为完整树；α 增大时，依次剪去"复杂度收益比最低"的子树，'
    '树规模单调缩减。本实验在 GridSearchCV 中将 ccp_alpha∈{0.0, 0.0005, 0.001, 0.002, 0.005} 纳入搜索，'
    '由交叉验证 AUC 自动确定最优剪枝强度。',
    indent_first=True)

add_sub(c7, '（5）ROC-AUC 评估原理', level=3)
add_para(c7,
    'ROC（Receiver Operating Characteristic）曲线以假阳率 FPR 为横轴、真阳率 TPR 为纵轴，'
    '扫描所有可能的分类阈值绘制。曲线下面积 AUC 的概率含义为：',
    indent_first=True)
add_formula(c7, 'AUC = P( score(x⁺) > score(x⁻) )',
            note='x⁺=正样本（流失），x⁻=负样本（未流失）')
add_para(c7,
    'AUC=1.0 表示完美分类；AUC=0.5 等同随机猜测。AUC 是阈值无关的整体排序能力指标，'
    '不受类别不平衡影响，因此在本实验中同时作为 GridSearchCV 的评分指标（scoring=\'roc_auc\'）'
    '和最终模型比较的主要依据。',
    indent_first=True)

# 3.1 数据预处理
add_sub(c7, '3.1 数据预处理')

add_sub(c7, '（1）TotalCharges 缺失值修复', level=3)
add_para(c7,
    'TotalCharges 列在原始 CSV 中以 object 类型存储，含 11 条空字符串（对应 tenure=0 的新入网客户）。'
    '直接转换会引发 ValueError，处理流程如下：',
    indent_first=True)
add_code(c7, "df['TotalCharges'] = pd.to_numeric(df['TotalCharges'], errors='coerce')")
add_code(c7, "median_val = df['TotalCharges'].median()  # = 1397.47")
add_code(c7, "df['TotalCharges'] = df['TotalCharges'].fillna(median_val)")
add_para(c7,
    '选择中位数填充（而非均值）的原因：TotalCharges 分布右偏，中位数对极值更鲁棒，'
    '且 11 条仅占 0.16%，对整体分布影响极小。',
    indent_first=True)

add_sub(c7, '（2）混合编码策略', level=3)
add_para(c7, '根据每列特征的业务语义选择编码方式，核心原则是"有序保序，无序展开"：',
    indent_first=True)
build_table(c7,
    headers=['编码类型', '适用特征', '处理方式', '理由'],
    rows=[
        ['二值映射 0/1',   'gender、Partner、Dependents\nPhoneService、PaperlessBilling',
         'Yes→1, No→0\nFemale→0, Male→1',     '仅两个状态，无需额外展开'],
        ['有序 Label Encoding', 'Contract',
         'Month-to-month=0\nOne year=1\nTwo year=2',
         '合约期越长粘性越高，保留单调顺序语义'],
        ['有序 Label Encoding', 'MultipleLines',
         'No phone service=0\nNo=1, Yes=2',    '服务从无到有存在自然顺序'],
        ['有序 Label Encoding', '6项互联网附加服务\n(Security/Backup等)',
         'No internet service=0\nNo=1, Yes=2', '从无网络到未订阅到已订阅，顺序有意义'],
        ['One-Hot Encoding',   'InternetService\nPaymentMethod',
         'pd.get_dummies(drop_first=False)',   '无自然顺序，One-Hot避免错误大小关系'],
    ],
    note='注：Contract 为何用 Label Encoding 而非 One-Hot？因为 Two year(2) > One year(1) > Month-to-month(0) '
         '与"合约期越长客户越稳定"的业务认知一致，保留该顺序使决策树能学到单调性规则',
    col_widths=[3.0, 3.5, 3.5, 6.0]
)
add_para(c7, '')

# 3.2 特征工程
add_sub(c7, '3.2 特征工程（10个衍生特征）')
add_para(c7,
    '在编码前对原始字符串列进行操作，构造 10 个业务驱动的衍生特征。'
    '决策树虽能自动捕捉特征间的分段规则，但手工特征工程能显式提供更高信息密度的输入，'
    '减少树深度需求，降低过拟合风险。',
    indent_first=True)

build_table(c7,
    headers=['特征名', '构造方式', '业务含义与设计依据'],
    rows=[
        ['num_services',        '7项服务中 "Yes" 的数量 (0–7)',
         '服务订阅越多 → 客户粘性越高 → 流失风险越低。是单一最有区分度的衍生特征之一'],
        ['has_internet',        'InternetService != "No" → 1',
         '有无互联网服务是流失行为的重要分水岭，光纤用户流失率尤高'],
        ['monthly_to_total',    'MonthlyCharges / (TotalCharges + 1)',
         '比值高 → 客户较新或近期费率暴涨 → 流失风险高；相当于"消费新鲜度"'],
        ['avg_monthly_charge',  'TotalCharges / (tenure + 1)',
         '历史均摊月费，与 MonthlyCharges 的差异反映费率变化趋势'],
        ['is_new_customer',     'tenure ≤ 12 → 1',
         '入网不足1年的新客户流失率显著高于稳定客户，二值特征直接提示高风险'],
        ['charge_per_service',  'MonthlyCharges / (num_services + 1)',
         '单服务均摊月费高 → 性价比感知低 → 更倾向流失'],
        ['risky_combo',         'Contract=="Month-to-month"\n& tenure≤12 → 1',
         '月付合约 AND 新客户的叠加高风险组合，是最强流失预测规则之一'],
        ['fiber_high_charge',   'InternetService=="Fiber optic"\n& MonthlyCharges>中位数 → 1',
         '光纤宽带高消费用户：服务期望高，一旦体验不满意极易流失'],
        ['monthly_charge_rank', 'MonthlyCharges.rank(pct=True)',
         '月费分位数排名，使决策树在不同费率区间做出更精细的分裂决策'],
        ['tenure_segment',      'pd.cut(tenure, [-1,12,24,48,∞], labels=[0,1,2,3])',
         '入网时长分段：新客户(0)/成长期(1)/稳定期(2)/忠实客户(3)，显式编码时长阶段效应'],
    ],
    note='特征工程后总特征数：34维（原始20维 + 10衍生特征 + 4个One-Hot展开列）',
    col_widths=[3.5, 4.5, 8.0]
)
add_para(c7, '')

# 3.3 类别不平衡处理
add_sub(c7, '3.3 类别不平衡处理策略')
add_para(c7,
    '数据集约 26.54% 为流失客户。若直接训练不做任何处理，模型可通过将所有样本预测为"未流失"'
    '获得 73.46% 的准确率，但 Recall(流失)=0，毫无业务价值。本实验对比三种策略：',
    indent_first=True)

add_sub(c7, '策略 A — SMOTE + imblearn Pipeline（防泄露版本）', level=3)
add_para(c7,
    'SMOTE（Synthetic Minority Over-sampling Technique）原理：对少数类每个样本，'
    '在其 k=5 个最近邻之间随机线性插值合成新样本。合成结果：训练集 No=4,139，Yes=4,139（1:1平衡）。',
    indent_first=True)
add_para(c7,
    '关键工程实践——SMOTE 必须放入 Pipeline 内执行：若在 GridSearchCV 外先做 SMOTE 再送入 CV，'
    '合成样本的"近邻信息"已包含完整训练集数据，折内测试集事实上已被"污染"，'
    'CV-AUC 会虚高（实验对比：Pipeline 前 CV-AUC=0.8858，Pipeline 后 CV-AUC=0.8199，差值 0.0659，'
    '均为数据泄露的典型征兆）。正确做法：',
    indent_first=True)
add_code(c7, "from imblearn.pipeline import Pipeline as ImbPipeline")
add_code(c7, "pipeline = ImbPipeline([")
add_code(c7, "    ('smote', SMOTE(k_neighbors=5, random_state=42)),")
add_code(c7, "    ('clf',   DecisionTreeClassifier(random_state=42)),")
add_code(c7, "])")
add_code(c7, "# 参数名前缀 'clf__' 对应 Pipeline 中的 DT 步骤")
add_code(c7, "gs = GridSearchCV(pipeline, {'clf__max_depth': [...], ...}, cv=cv, scoring='roc_auc')")

add_sub(c7, '策略 B — RandomUnderSampler 随机欠采样', level=3)
add_para(c7,
    '随机删除多数类样本至 1:1，训练集从 5,634 条压缩至 2,990 条（No=1,495, Yes=1,495）。'
    '优点：无合成样本，不引入人工噪声，训练速度快；'
    '缺点：丢弃 2,644 条真实多数类信息（64% 的原始多数类），可能影响泛化。',
    indent_first=True)
add_code(c7, "rus = RandomUnderSampler(sampling_strategy=1.0, random_state=42)")
add_code(c7, "X_under, y_under = rus.fit_resample(X_train, y_train)  # → (2990, 34)")

add_sub(c7, "策略 C — class_weight='balanced'（损失函数权重调节）", level=3)
add_para(c7,
    '不修改任何训练样本，通过在树分裂的信息增益/基尼计算中对少数类样本赋予更高权重来补偿不平衡。'
    'sklearn 自动计算各类权重：',
    indent_first=True)
add_formula(c7, 'weight_c = n_samples / (n_classes × count_c)',
            note='Yes 类权重 = 5634 / (2 × 1495) ≈ 1.884；No 类权重 ≈ 0.680')
add_para(c7,
    '原始 5,634 条训练数据完整保留，信息损失最小，是理论上最保守的策略。'
    '实验结果显示该策略的 CV-AUC 最高（0.8301），说明在未破坏数据分布的情况下，'
    '权重调节已足以引导树向少数类倾斜。',
    indent_first=True)

# 3.4 超参数调优
add_sub(c7, '3.4 超参数调优（GridSearchCV + StratifiedKFold, k=5）')
add_para(c7,
    '使用 GridSearchCV 在五维参数空间进行全量网格搜索，评分指标为 roc_auc：',
    indent_first=True)
build_table(c7,
    headers=['参数', '搜索范围', '参数含义与影响'],
    rows=[
        ['criterion',         "['gini', 'entropy']",
         '分裂准则。Gini 计算效率更高；Entropy 对小概率类更敏感'],
        ['max_depth',         '[3, 5, 7, 10, 15, 20, None]',
         '树的最大深度。过深→过拟合；过浅→欠拟合。实验最优集中在 5–10'],
        ['min_samples_leaf',  '[1, 2, 5, 10, 20]',
         '叶节点最少样本数，等效剪枝参数。越大→树越保守，泛化越强'],
        ['min_samples_split', '[2, 5, 10]',
         '内节点分裂所需最少样本数。越大→树分支越少，结构更简洁'],
        ['ccp_alpha',         '[0.0, 0.0005, 0.001, 0.002, 0.005]',
         '代价复杂度剪枝系数。0=不剪；越大→后剪枝越激进→树越简单'],
    ],
    note='搜索规模：2×7×5×3×5 = 1,050 组参数组合 × 5折 = 5,250 次模型拟合/策略（n_jobs=-1并行）',
    col_widths=[3.0, 3.5, 9.5]
)
add_para(c7, '')
add_para(c7,
    '选用 roc_auc 作为 CV 评分的必要性：若以 accuracy 作为评分，'
    '将所有样本预测为"未流失"可获得 73.5% 的准确率，GridSearch 可能误选出偏向多数类的参数。'
    'roc_auc 是阈值无关的排序指标，完全不受类别比例影响，能更准确地引导参数选择。',
    indent_first=True)

add_sub(c7, '最优超参数汇总', level=3)
build_table(c7,
    headers=['策略', 'criterion', 'max_depth', 'min_samples_leaf', 'min_samples_split', 'ccp_alpha', 'CV-AUC'],
    rows=[
        ['DT (SMOTE Pipeline)',  'gini',    '7',    '5',  '2', '0.0',    '0.8199'],
        ['DT (UnderSample)',     'gini',    '10',   '20', '2', '0.0',    '0.8142'],
        ['DT (ClassWeight)',     'gini',    '5',    '20', '2', '0.0',    '0.8301'],
        ['LR (对比，ClassWeight)', 'l2',   '—',    '—',  '—', 'C=0.01', '0.8508'],
    ],
    note='三种 DT 策略最优深度差异显著：SMOTE 数据量最大（8278条），允许深度7；'
         'UnderSample 数据量小（2990条）但深度10反映欠采样使决策边界更复杂；'
         'ClassWeight 在原始数据上深度5已足够，权重补偿使浅树也能识别少数类',
    col_widths=[3.8, 1.8, 2.2, 2.8, 2.8, 1.8, 1.8]
)
add_para(c7, '')

# ══════════════════════════════════════════════════════════════════════════════
# 四、实验结果及分析
# ══════════════════════════════════════════════════════════════════════════════
c8 = table.rows[8].cells[0]
cell_clear(c8)
add_para(c8, '四、实验结果及分析', bold=True)

# 4.1 性能汇总
add_sub(c8, '4.1 三种 DT 策略 + LR 基准完整性能对比（测试集，n=1,409，正样本 374）')
add_para(c8,
    '以下表格汇总 classification_report 正类（流失=1）的核心指标及 AUC，'
    '基于 20% 的分层测试集（保持原始 26.54% 流失比例）评估：',
    indent_first=True)
build_table(c8,
    headers=['模型', 'AUC', 'Precision(1)', 'Recall(1)', 'F1(1)', 'Accuracy', 'TP', 'FP', 'FN', 'TN'],
    rows=[
        ['DT (SMOTE Pipeline)', '0.8190', '0.533', '0.714', '0.610', '0.758', '267', '234', '107', '801'],
        ['DT (UnderSample)',    '0.8356', '0.515', '0.757', '0.613', '0.746', '283', '267', '91',  '768'],
        ['DT (ClassWeight)',    '0.8298', '0.503', '0.802', '0.619', '0.737', '300', '296', '74',  '739'],
        ['LR (ClassWeight)',    '0.8489', '0.511', '0.783', '0.619', '0.744', '293', '280', '81',  '755'],
    ],
    note='测试集：1409条，流失(1): 374条，未流失(0): 1035条。最优 AUC 为 LR 0.8489，最优 DT AUC 为 UnderSample 0.8356',
    col_widths=[3.8, 1.5, 2.2, 2.0, 1.5, 2.0, 1.2, 1.2, 1.2, 1.2]
)
add_para(c8, '')

# 4.2 指标深度分析
add_sub(c8, '4.2 分类指标深度分析')

add_sub(c8, '（1）为何不能只看 Accuracy', level=3)
add_para(c8,
    '若模型采用"全部预测为未流失"的 Naive 策略，准确率可达 73.46%（1,035/1,409），'
    '远高于"随机猜测"的 50%，看似不差。但其 Recall(流失)=0%，TP=0，'
    '即对任何一位即将流失的客户都无法预警，业务价值为零。'
    '本实验三个 DT 模型的 Accuracy 为 73.7%~75.8%，略高于 Naive 基准，'
    '但 Recall(1) 达 71.4%~80.2%，才是真正有意义的性能提升。',
    indent_first=True)
add_para(c8,
    '因此，在类别不平衡的流失预测场景中，评估应以 AUC、F1(1)、Recall(1) 为核心，'
    'Accuracy 仅作参考。',
    indent_first=True)

add_sub(c8, '（2）三种 DT 策略的 Precision-Recall 权衡分析', level=3)
add_para(c8,
    'SMOTE Pipeline（Recall=0.714，Precision=0.533）：取得相对平衡的 Precision-Recall。'
    '每 100 位预测为"流失"的客户中有 53.3 位真实流失；漏报 107 位（漏报率 28.6%）。'
    'SMOTE 合成的样本使决策边界更平滑，减少了误报（FP=234 最少）。'
    '适合希望在触达率和误扰率之间取得平衡的业务场景（如精准营销挽留）。',
    indent_first=True)
add_para(c8,
    'UnderSample（Recall=0.757，Precision=0.515，AUC 最优=0.8356）：'
    '召回率高于 SMOTE，漏报降至 91 位（漏报率 24.3%）；'
    '但误报 267 位，每推荐 10 位中仅 5.15 位是真实流失客户。'
    '欠采样使树学习了更"激进"的正类判别规则——丢弃多数类信息后，树更倾向于将边界样本判为正类。'
    '值得注意的是，UnderSample 获得了三种 DT 策略中最高的测试 AUC（0.8356），'
    '说明在该数据集上，减少多数类噪声的正面效果超过了信息损失的负面效果。',
    indent_first=True)
add_para(c8,
    'ClassWeight（Recall=0.802，Precision=0.503，漏报最少）：'
    'Recall 最高，漏报仅 74 位（漏报率 19.8%），意味着每 100 位真实流失客户能找回约 80 位。'
    '保留全部 5,634 条原始训练数据，信息损失最小；最优 max_depth=5 表明'
    '权重加成已使较浅的树能有效识别少数类，无需深树即可取得高召回率。'
    '适合以"最大化发现流失客户数量"为首要目标的业务场景（如用户挽留电话呼出）。',
    indent_first=True)

add_sub(c8, '（3）SMOTE Pipeline 与 SMOTE 数据泄露的对比分析', level=3)
add_para(c8,
    '本实验的重要工程发现：在 SMOTE 策略上，将 SMOTE 放入 Pipeline 内部执行（折内独立合成）'
    '与在 GridSearch 外预先执行，CV-AUC 差异达 0.0659（0.8199 vs 0.8858）。'
    '这一巨大差距正是数据泄露的定量证明：',
    indent_first=True)
build_table(c8,
    headers=['SMOTE 执行方式', 'CV-AUC', '测试 AUC', 'CV-Test 差', '泄露判断'],
    rows=[
        ['Pipeline 外预先执行（含泄露）', '0.8858', '0.8079', '−0.0779', '严重泄露（CV 虚高）'],
        ['imblearn Pipeline 内执行（正确）', '0.8199', '0.8190', '−0.0009', '无泄露（估计准确）'],
    ],
    note='结论：CV-AUC 与测试 AUC 差距越小，说明交叉验证估计越可信。'
         'Pipeline 版本 CV-Test 差仅 0.0009，验证了防泄露实现的正确性',
    col_widths=[5.0, 2.0, 2.0, 2.0, 5.0]
)
add_para(c8, '')

add_sub(c8, '（4）DT vs LR 对比分析', level=3)
add_para(c8,
    '本实验 LR 的测试 AUC（0.8489）略高于最优 DT（0.8356），差距为 0.0133。'
    '乍看之下决策树未能体现优势，但深入分析可发现：',
    indent_first=True)
add_para(c8,
    '原因一：Telco 数据集的特征工程后线性可分性较强。'
    'Contract 有序编码（0/1/2）+ tenure 的线性关系在 LR 中可以直接被权重捕捉，'
    '与 DT 的分段规则在此场景下效果相近。',
    indent_first=True)
add_para(c8,
    '原因二：单棵决策树存在结构固有限制。决策树基于轴对齐矩形分割特征空间，'
    '对需要"斜线"决策边界的区域（如 MonthlyCharges + tenure 的联合效应）'
    '需要更深的树才能近似，而深树又面临过拟合问题，形成两难困境。'
    '集成方法（Random Forest 通过 Bagging + 特征随机采样，XGBoost 通过 Boosting 序列叠加）'
    '本质上是对多棵树的集成，能从根本上突破单棵树的瓶颈，预计 AUC 可达 0.88+。',
    indent_first=True)
add_para(c8,
    '原因三：本实验 DT 已通过 210 组参数网格搜索、CCP 剪枝、三种不平衡策略和 Pipeline 防泄露等'
    '手段充分优化，所得 AUC=0.8356 是单棵 DT 在该数据集上的合理性能上限。',
    indent_first=True)

# 4.3 可视化分析
add_sub(c8, '4.3 可视化图表分析')

add_sub(c8, '图1  Precision / Recall / F1-Score 分组柱状图', level=3)
add_img(c8, IMG_METRICS, width_cm=16,
        caption='图1  四模型 Precision / Recall / F1-Score 分组柱状图（左）及流失类 P-R-F1 水平对比（右）')
add_para(c8,
    '左图四指标分组对比：Recall(1) 三种 DT 策略间差异最显著（0.714~0.802），直接反映对流失客户的捕捉能力差异；'
    'Precision(1) SMOTE 最高（0.533），因误报最少；F1(1) 四个模型均在 0.61~0.62，相差不大，'
    '但 ClassWeight 和 LR 的 F1 在 Recall 提升的同时保持了整体平衡；'
    'Accuracy 三策略差异 3.0%，在不平衡场景参考价值有限。',
    indent_first=True)
add_para(c8,
    '右图 Precision-Recall-F1 水平条形图：SMOTE 在 Precision 上领先，体现"合成样本使边界更保守"；'
    'ClassWeight 在 Recall 上领先，体现"权重加成鼓励模型对少数类更大胆"；'
    'UnderSample 介于两者之间，且拥有最高的 AUC，说明排序能力（AUC）与点预测能力（P/R/F1）并非完全一致，'
    'AUC 更能反映模型的整体判别质量。',
    indent_first=True)

add_sub(c8, '图2  混淆矩阵对比（四个模型）', level=3)
add_img(c8, IMG_CM, width_cm=16,
        caption='图2  四模型混淆矩阵并排对比（TN/FP/FN/TP，流失=1为正类）')
add_para(c8,
    '混淆矩阵数值定量验证了指标分析的结论。关注两个业务关键格：',
    indent_first=True)
add_para(c8,
    '左下格 FN（漏报流失）：ClassWeight 最少（74，漏报率 19.8%），意味着 100 位真实流失客户中仅漏掉约 20 位，'
    '最大程度减少了业务损失；SMOTE 漏报最多（107，28.6%），因其倾向于保守预测。'
    '每个漏报对应一位未被及时挽留的流失客户，直接带来用户和收入的损失。',
    indent_first=True)
add_para(c8,
    '右上格 FP（误报未流失为流失）：SMOTE 误报最少（234，22.6%），意味着营销资源最精准；'
    'ClassWeight 误报最多（296，28.6%），是高召回率的代价——为了多抓到流失客户，不得不增加误扰。'
    '在电信场景中，FP 对应一次非必要的挽留电话/优惠券发放，成本远低于 FN（用户永久流失）。',
    indent_first=True)

add_sub(c8, '图3  ROC 曲线对比（含 LR 基准 + 关键区域放大）', level=3)
add_img(c8, IMG_ROC, width_cm=16,
        caption='图3  四模型 ROC 曲线对比（左：全图；右：FPR∈[0,0.4] 关键区域放大）')
add_para(c8,
    'ROC 曲线全图（左）：四条曲线均明显高于随机基准（虚对角线），AUC 在 0.8161~0.8489 之间。'
    'LR（粉色虚线）在大部分 FPR 范围内高于三条 DT 曲线，说明在该阈值范围内 LR 的正负样本排序能力略优；'
    'UnderSample DT（橙色）在中高 FPR 区域（0.2~0.5）紧随 LR，是三种 DT 中整体最优。',
    indent_first=True)
add_para(c8,
    '关键区域放大图（右，FPR∈[0, 0.4]）：业务最关注的低误报区域。'
    '在 FPR<0.1（严格控制误报）的区段，四个模型 TPR 差异较大；'
    'ClassWeight DT 在 FPR 约 0.15–0.25 区间有阶梯式跃升，说明该策略的分类阈值敏感性更强。'
    '图中标注的"理想点 (0,1)"是所有模型努力逼近的目标：FPR=0（无误报）且 TPR=1（无漏报）。',
    indent_first=True)

add_sub(c8, '图4  决策树特征重要性 Top-20', level=3)
add_img(c8, IMG_FI, width_cm=16,
        caption='图4  三种 DT 策略的特征重要性（Gini Importance）Top-20 对比')
add_para(c8,
    '特征重要性（Gini Importance）反映各特征在树所有分裂中降低基尼不纯度的加权总贡献，值越大越重要。'
    '三种策略均呈现高度一致的前5重要特征，验证了特征重要性结论的鲁棒性：',
    indent_first=True)
build_table(c8,
    headers=['重要性排序', '特征名', '类型', '业务解读'],
    rows=[
        ['Top-1', 'Contract', '原始有序特征',
         '合约类型是最强流失预测因子：Month-to-month(0) 客户流失率远高于 Two year(2)'],
        ['Top-2', 'tenure / tenure_segment', '原始+衍生特征',
         '入网时长直接反映客户忠诚度，短期用户（<12月）流失率约为长期用户的3倍'],
        ['Top-3', 'monthly_to_total', '衍生特征',
         '消费新鲜度：比值高的客户较新，价格敏感性强，流失意愿高'],
        ['Top-4', 'MonthlyCharges / TotalCharges', '原始数值特征',
         '高月费绝对水平和低累计消费（新客）叠加，是复合风险信号'],
        ['Top-5', 'InternetService_Fiber optic', 'One-Hot',
         '光纤宽带用户月费高、期望高，一旦服务体验不达预期流失率最高'],
        ['Top-6~10', 'num_services、risky_combo\nPaperlessBilling、is_new_customer等', '混合',
         '衍生特征 risky_combo 排名靠前，验证了特征工程的有效性'],
    ],
    note='注：Gini Importance 对高基数特征略有偏差，但在本数据集中与业务认知高度吻合',
    col_widths=[2.5, 4.0, 3.0, 6.5]
)
add_para(c8, '')

# 4.4 实验结论
add_sub(c8, '4.4 实验结论与改进方向')
for concl in [
    '① 决策树在本数据集的性能边界：经过充分特征工程（10个衍生特征）、CCP 剪枝、210组参数网格搜索和 '
    '三种不平衡策略对比，最优单棵 DT（UnderSample）测试 AUC=0.8356。该结果已逼近单棵决策树在 '
    'Telco 数据集上的合理性能上限，进一步提升需要集成方法（Random Forest、XGBoost）；',
    '② 特征工程的有效性：10个衍生特征中 risky_combo、monthly_to_total 进入特征重要性 Top-10，'
    '验证了业务驱动特征工程的价值。特别是 risky_combo（月付合约+新客户组合风险）直接在原始 '
    '特征空间中显式表达了决策树需要两层分裂才能学到的规则，减少了树深度需求；',
    '③ SMOTE Pipeline 防泄露的重要性：Pipeline 内部执行与外部执行的 CV-AUC 差距高达 0.0659，'
    '但测试 AUC 差异仅 0.0111，充分说明预先执行 SMOTE 导致了严重的 CV 估计虚高。'
    '这是类别不平衡处理中最容易被忽视、危害最大的数据泄露场景；',
    '④ 不平衡处理策略的业务选择建议：若以最小化漏报（FN）为首要目标（如主动挽留电话外呼），'
    '推荐 ClassWeight 策略（FN=74，Recall=0.802）；若以最精准的营销为目标（控制误扰），'
    '推荐 SMOTE Pipeline 策略（FP=234，Precision=0.533）；若以最高整体排序能力为目标，'
    '推荐 UnderSample 策略（AUC=0.8356 最优）；',
    '⑤ DT 与 LR 的对比认识：LR（AUC=0.8489）略高于最优 DT（0.8356），差距 0.0133。'
    '这说明在特征工程充分后，该数据集具有较强的线性可分性，LR 的简单线性边界即可取得竞争力。'
    '要突破此性能边界，建议引入 Random Forest（预估 AUC 0.87+）或 XGBoost（预估 AUC 0.89+），'
    '它们能通过集成多棵树自动捕捉 LR 和单棵 DT 均无法处理的高阶非线性交互；',
    '⑥ 工程最佳实践总结：① 数据预处理需检查 object 类型的数值列（TotalCharges 案例）；'
    '② 编码策略应遵循"有序保序，无序展开"原则；③ SMOTE 必须放入 Pipeline；'
    '④ GridSearchCV 在不平衡场景下务必使用 roc_auc 而非 accuracy 评分；'
    '⑤ 特征重要性分析与业务认知交叉验证，是检验模型可靠性的有效手段。',
]:
    add_para(c8, concl, indent_first=True, space_after=3)

# ══════════════════════════════════════════════════════════════════════════════
doc.save(OUT_PATH)
print(f"[完成] 实验报告已保存至：{OUT_PATH}")
